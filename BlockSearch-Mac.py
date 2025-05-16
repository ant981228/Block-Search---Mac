import docx
import json
import os
import re
import subprocess
import sys
import time
import tkinter as tk
import threading
import tempfile
import zipfile
from docx.document import Document
from docx.text.paragraph import Paragraph
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
from tempfile import TemporaryDirectory
from tkinter import filedialog, messagebox, ttk

def get_pandoc_path():
    """Get the path to the bundled Pandoc executable in a Mac app."""
    import os
    import sys
    import subprocess
    
    # When run as a PyInstaller bundle
    if getattr(sys, 'frozen', False):
        if sys.platform == 'darwin':  # macOS
            # Look for pandoc in the specific subdirectory we defined in the spec file
            bundle_dir = os.path.dirname(sys.executable)
            
            # Primary search locations - check in the pandoc-bin directory we specified
            pandoc_bin_dir = os.path.join(os.path.dirname(bundle_dir), 'Resources', 'pandoc-bin')
            pandoc_path = os.path.join(pandoc_bin_dir, 'pandoc')
            
            if os.path.exists(pandoc_path) and os.access(pandoc_path, os.X_OK):
                print(f"Found bundled pandoc at: {pandoc_path}")
                return pandoc_path
                
            # Secondary search locations - other common places PyInstaller might put it
            possible_paths = [
                os.path.join(bundle_dir, 'pandoc-bin', 'pandoc'),  # In MacOS/pandoc-bin directory
                os.path.join(bundle_dir, 'pandoc'),                # Directly in MacOS directory
                os.path.join(os.path.dirname(bundle_dir), 'Resources', 'pandoc'),  # In Resources
                os.path.join(os.path.dirname(os.path.dirname(bundle_dir)), 'pandoc-bin', 'pandoc'),  # App root
            ]
            
            # Try each possible path
            for path in possible_paths:
                if os.path.exists(path) and os.access(path, os.X_OK):
                    print(f"Found bundled pandoc at: {path}")
                    return path
            
            # Print debugging information if we can't find pandoc
            print("\n--- PANDOC DEBUGGING INFO ---")
            print(f"Bundle directory: {bundle_dir}")
            print("Checking possible pandoc locations:")
            
            # Check all possible locations and print results
            search_paths = [pandoc_bin_dir] + [os.path.dirname(p) for p in possible_paths]
            for path in sorted(set(search_paths)):  # Use set to avoid duplicates
                exists = os.path.exists(path)
                print(f"  Directory: {path}")
                print(f"    Exists: {exists}")
                
                if exists:
                    print("    Contents:")
                    try:
                        files = os.listdir(path)
                        for file in files:
                            file_path = os.path.join(path, file)
                            is_exec = os.access(file_path, os.X_OK)
                            print(f"      {file} (executable: {is_exec})")
                    except Exception as e:
                        print(f"      Error listing directory: {e}")
            
            # Last resort - try to find pandoc on the system path
            try:
                result = subprocess.run(['which', 'pandoc'], 
                                      stdout=subprocess.PIPE, 
                                      stderr=subprocess.PIPE,
                                      text=True,
                                      check=False)
                
                if result.returncode == 0 and result.stdout.strip():
                    system_path = result.stdout.strip()
                    print(f"Found system pandoc at: {system_path}")
                    
                    # Verify it's executable
                    if os.access(system_path, os.X_OK):
                        print("System pandoc is executable, using it")
                        return system_path
            except Exception as e:
                print(f"Error checking system pandoc: {e}")
                
            print("--- END PANDOC DEBUGGING INFO ---\n")
            
            # If we get here, we couldn't find pandoc - display a clear error
            print("ERROR: Pandoc not found in the bundled application or system PATH.")
            print("Please make sure pandoc is installed and accessible.")
            
            # Return the regular command as a last resort, but likely to fail
            return "pandoc"
        else:
            # Windows or other platform
            return os.path.join(sys._MEIPASS, 'pandoc-bin', 'pandoc.exe' if sys.platform == 'win32' else 'pandoc')
    
    # In development mode - try to use system pandoc
    try:
        # Try to find pandoc in PATH
        result = subprocess.run(['which', 'pandoc'] if sys.platform != 'win32' else ['where', 'pandoc'],
                              stdout=subprocess.PIPE, 
                              stderr=subprocess.PIPE,
                              text=True,
                              check=False)
        
        if result.returncode == 0 and result.stdout.strip():
            pandoc_path = result.stdout.strip().split('\n')[0]  # Take first result if multiple
            print(f"Development mode: Using pandoc from PATH: {pandoc_path}")
            return pandoc_path
    except Exception as e:
        print(f"Error finding pandoc in development mode: {e}")
    
    # Default fallback
    print("Using default 'pandoc' command (not found in PATH)")
    return "pandoc"

class JSONToWordConverter:
    """Elegant macOS-style application for converting JSON files to Word and copying to clipboard."""
    
    # Constants
    CONFIG_FILE = os.path.expanduser("~/Library/Application Support/JSONToWordConverter/config.json")
    DEFAULT_TEMPLATE_NAME = "template_doc.docx"
    
    def __init__(self, root):
        self.root = root
        self.root.title("Block Search")
        
        # Variables
        self.directory_var = tk.StringVar()
        self.search_var = tk.StringVar()
        self.template_var = tk.StringVar()
        self.status_var = tk.StringVar(value="Ready")
        self.reverse_sort_var = tk.BooleanVar(value=False)
        self.target_document_var = tk.StringVar(value="Copy to clipboard only")
        self.open_documents = ["Copy to clipboard only"]
        
        # Current sort settings
        self.current_sort = {"column": "filename", "reverse": False}
        
        # Add dark mode variable
        self.dark_mode = tk.BooleanVar(value=False)
        
        # Add trace to update UI when dark mode changes
        self.dark_mode.trace_add("write", self.update_theme)

        # Setup UI with macOS styling
        self.setup_ui()
        
        # Setup menu
        self.setup_menu()

        # Check for permissions
        self.root.after(1000, self.check_accessibility_permissions)
        
        # Load configuration and initialize
        self.load_config()

        # Bind search updates
        self.search_var.trace_add("write", self.search_files)
        
        # Bind keyboard shortcuts
        self.root.bind("<Command-r>", self.refresh_word_documents)
        
        # Standard search focus shortcuts
        if sys.platform == 'darwin':  # macOS
            self.root.bind('<Command-f>', self.focus_search_field)
        else:  # Windows/Linux
            self.root.bind('<Control-f>', self.focus_search_field)
            
        # Add Control+Shift+Space shortcut (works on all platforms)
        self.root.bind('<Control-Shift-space>', self.focus_search_field)

    def setup_ui(self):
        """Create the user interface with macOS styling."""
        # Configure ttk style
        self.configure_styles()
        
        # Create main frame
        main_frame = ttk.Frame(self.root, style="Main.TFrame")
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # Header
        header_frame = ttk.Frame(main_frame, style="Main.TFrame")
        header_frame.pack(fill=tk.X, pady=(0, 15))
        
        header_label = ttk.Label(header_frame, text="Block Search", 
                                style="Header.TLabel")
        header_label.pack(side=tk.LEFT)
        
        # Directory section
        dir_frame = ttk.Frame(main_frame, style="Main.TFrame")
        dir_frame.pack(fill=tk.X, pady=5)
        
        dir_label = ttk.Label(dir_frame, text="Directory:", 
                             style="Label.TLabel")
        dir_label.pack(side=tk.LEFT, padx=(0, 10))
        
        dir_entry = ttk.Entry(dir_frame, textvariable=self.directory_var, 
                             style="Entry.TEntry", width=50)
        dir_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 10))
        
        browse_button = ttk.Button(dir_frame, text="Browse", 
                                  command=self.browse_directory,
                                  style="Button.TButton")
        browse_button.pack(side=tk.LEFT)
        
        # Template section
        template_frame = ttk.Frame(main_frame, style="Main.TFrame")
        template_frame.pack(fill=tk.X, pady=5)
        
        template_label = ttk.Label(template_frame, text="Template:", 
                                  style="Label.TLabel")
        template_label.pack(side=tk.LEFT, padx=(0, 10))
        
        template_entry = ttk.Entry(template_frame, textvariable=self.template_var, 
                                  style="Entry.TEntry", width=50)
        template_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 10))
        
        template_button = ttk.Button(template_frame, text="Browse", 
                                    command=self.browse_template,
                                    style="Button.TButton")
        template_button.pack(side=tk.LEFT)
        
        # Search section
        search_frame = ttk.Frame(main_frame, style="Main.TFrame")
        search_frame.pack(fill=tk.X, pady=(15, 5))
        
        search_label = ttk.Label(search_frame, text="Search:", 
                                style="Label.TLabel")
        search_label.pack(side=tk.LEFT, padx=(0, 10))
        
        search_entry = ttk.Entry(search_frame, textvariable=self.search_var, 
                                style="Entry.TEntry")
        search_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        # Store a reference to the search entry for focus management
        self.search_entry = search_entry
        
        # Add keyboard bindings for search entry
        self.search_entry.bind("<Down>", self.handle_search_down_key)
        self.search_entry.bind("<Up>", self.handle_search_up_key)
        self.search_entry.bind("<Return>", self.handle_search_enter_key)
        
        # File list with scrollbar
        list_frame = ttk.Frame(main_frame, style="Main.TFrame")
        list_frame.pack(fill=tk.BOTH, expand=True, pady=10)
        
        # Create horizontal and vertical scrollbars
        v_scrollbar = ttk.Scrollbar(list_frame, orient="vertical")
        v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        h_scrollbar = ttk.Scrollbar(list_frame, orient="horizontal")
        h_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)
        
        # Create treeview with multiple columns
        self.tree = ttk.Treeview(list_frame, 
                                columns=("filename", "relpath", "lastopened"), 
                                show="headings", 
                                yscrollcommand=v_scrollbar.set,
                                xscrollcommand=h_scrollbar.set)
        
        # Configure the columns
        self.tree.heading("filename", text="File Name", command=lambda: self.sort_treeview("filename", False))
        self.tree.heading("relpath", text="Relative Path", command=lambda: self.sort_treeview("relpath", False))
        self.tree.heading("lastopened", text="Last Opened", command=lambda: self.sort_treeview("lastopened", False))
        
        # Set column widths and anchors
        self.tree.column("filename", width=400, anchor=tk.W, stretch=True)
        self.tree.column("relpath", width=200, anchor=tk.W)
        self.tree.column("lastopened", width=150, anchor=tk.W)
        
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # Connect scrollbars to treeview
        v_scrollbar.config(command=self.tree.yview)
        h_scrollbar.config(command=self.tree.xview)
        
        # Add keyboard bindings for treeview navigation
        print("\n=== Setting up event bindings ===")
        print("Binding <Double-1> to on_file_select")  
        self.tree.bind("<Return>", self.on_file_select)
        self.tree.bind("<Double-1>", self.on_double_click)
        self.tree.bind("<Up>", lambda event: self.handle_tree_navigation(event, "up"))
        self.tree.bind("<Down>", lambda event: self.handle_tree_navigation(event, "down"))
        
        # Check for any Return/Enter bindings
        if hasattr(self, 'handle_search_enter_key'):
            print("Binding <Return> to handle_search_enter_key in search entry")
        if hasattr(self, 'handle_tree_navigation'):
            print("Binding Up/Down to handle_tree_navigation in treeview")
        
        # Any additional bindings that might trigger file selection
        for binding in self.tree.bind():
            print(f"Tree already has binding: {binding}")

        # Controls frame containing sort options and target selection
        controls_frame = ttk.Frame(main_frame, style="Main.TFrame")
        controls_frame.pack(fill=tk.X, pady=5)
        
        # Left side: Sort options
        sort_frame = ttk.Frame(controls_frame, style="Main.TFrame")
        sort_frame.pack(side=tk.LEFT, fill=tk.Y)
        
        sort_label = ttk.Label(sort_frame, text="Sort by:", style="Label.TLabel")
        sort_label.pack(side=tk.LEFT, padx=(0, 10))
        
        # Sort buttons
        name_sort_btn = ttk.Button(sort_frame, text="Name", 
                                  command=lambda: self.sort_treeview("filename", False),
                                  style="Button.TButton", width=10)
        name_sort_btn.pack(side=tk.LEFT, padx=2)
        
        path_sort_btn = ttk.Button(sort_frame, text="Path", 
                                  command=lambda: self.sort_treeview("relpath", False),
                                  style="Button.TButton", width=10)
        path_sort_btn.pack(side=tk.LEFT, padx=2)
        
        date_sort_btn = ttk.Button(sort_frame, text="Last Opened", 
                                  command=lambda: self.sort_treeview("lastopened", False),
                                  style="Button.TButton", width=12)
        date_sort_btn.pack(side=tk.LEFT, padx=2)
        
        # Reverse sort checkbox
        self.reverse_sort_var = tk.BooleanVar(value=False)
        reverse_check = ttk.Checkbutton(sort_frame, text="Reverse", 
                                       variable=self.reverse_sort_var,
                                       command=self.refresh_sort,
                                       style="Checkbox.TCheckbutton")
        reverse_check.pack(side=tk.LEFT, padx=(20, 0))
        
        # Right side: Target document selection
        target_frame = ttk.Frame(controls_frame, style="Main.TFrame")
        target_frame.pack(side=tk.RIGHT, fill=tk.Y)
        
        target_label = ttk.Label(target_frame, text="Target:", 
                               style="Label.TLabel")
        target_label.pack(side=tk.LEFT, padx=(20, 5))
        
        # Document selection dropdown
        self.target_dropdown = ttk.Combobox(target_frame, 
                                          textvariable=self.target_document_var,
                                          style="Combobox.TCombobox",
                                          state="readonly",
                                          width=20)
        self.target_dropdown['values'] = self.open_documents
        self.target_dropdown.pack(side=tk.LEFT, padx=(0, 5))
        
        # Refresh button with icon instead of text
        refresh_button = ttk.Button(target_frame, text="⟳", 
                                   command=self.refresh_word_documents,
                                   style="Button.TButton", width=2)
        refresh_button.pack(side=tk.LEFT)
        
        # Status bar
        status_frame = ttk.Frame(main_frame, style="Main.TFrame")
        status_frame.pack(fill=tk.X, pady=(10, 0))
        
        status_label = ttk.Label(status_frame, textvariable=self.status_var, 
                                style="Status.TLabel")
        status_label.pack(side=tk.LEFT)
        
        instruction_text = "Double-click a file to convert and copy to clipboard"
        instruction_label = ttk.Label(status_frame, text=instruction_text, 
                                     style="Instruction.TLabel")
        instruction_label.pack(side=tk.RIGHT)
        
        # Store current sort info
        self.current_sort = {"column": "filename", "reverse": False}

    def check_accessibility_permissions(self):
        """Check if the app has accessibility permissions and prompt if needed."""
        # AppleScript to check if we have permissions by attempting a simple keystroke operation
        check_script = '''
        try
            tell application "System Events"
                -- Just checking if we can access System Events
                -- We don't actually press any keys here
                set ui_enabled to UI elements enabled
                return true
            end tell
        on error
            return false
        end try
        '''
        
        try:
            process = subprocess.Popen(
                ['osascript', '-e', check_script],
                stdout=subprocess.PIPE, 
                stderr=subprocess.PIPE
            )
            stdout, stderr = process.communicate()
            
            result = stdout.decode().strip().lower()
            if result == "true":
                print("✓ Accessibility permissions are granted")
                return True
            else:
                print("✗ Accessibility permissions are needed")
                
                # Show a dialog explaining the permissions needed
                response = messagebox.askquestion(
                    "Permissions Required",
                    "BlockSearch needs accessibility permissions to paste content into Word.\n\n"
                    "Would you like to open System Preferences to grant these permissions?",
                    icon='info'
                )
                
                if response == 'yes':
                    # Open System Preferences to the right panel
                    self.open_accessibility_preferences()
                    
                return False
        except Exception as e:
            print(f"Error checking permissions: {e}")
            return False

    def open_accessibility_preferences(self):
        """Open System Preferences to the Accessibility permissions panel."""
        # Open System Preferences directly to Security -> Privacy -> Accessibility
        script = '''
        tell application "System Preferences"
            activate
            set current pane to pane id "com.apple.preference.security"
            delay 0.5
            tell application "System Events" to tell process "System Preferences"
                delay 0.5
                if exists tab group 1 of window "Security & Privacy" then
                    click radio button "Privacy" of tab group 1 of window "Security & Privacy"
                    delay 0.3
                    if exists row "Accessibility" of table 1 of scroll area 1 of tab group 1 of window "Security & Privacy" then
                        select row "Accessibility" of table 1 of scroll area 1 of tab group 1 of window "Security & Privacy"
                    end if
                end if
            end tell
        end tell
        '''
        
        try:
            subprocess.Popen(['osascript', '-e', script], 
                            stdout=subprocess.PIPE, 
                            stderr=subprocess.PIPE)
        except Exception as e:
            print(f"Error opening System Preferences: {e}")

    def setup_menu(self):
        """Setup application menu with settings and help."""
        # Create main menu bar
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)
        
        # File menu
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="File", menu=file_menu)
        file_menu.add_command(label="Browse Directory...", command=self.browse_directory)
        file_menu.add_command(label="Browse Template...", command=self.browse_template)
        file_menu.add_command(label="Refresh Word Documents", command=self.refresh_word_documents)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.root.quit)
        
        # Edit menu
        edit_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Edit", menu=edit_menu)
        
        # Add focus search command with all shortcuts
        if sys.platform == 'darwin':  # macOS
            shortcut_display = "⌘F or ⌃⇧Space"
        else:  # Windows/Linux
            shortcut_display = "Ctrl+F or Ctrl+Shift+Space"
            
        edit_menu.add_command(
            label="Focus Search", 
            command=self.focus_search_field,
            accelerator=shortcut_display
        )
        
        # Document Tools menu
        document_tools_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Document Tools", menu=document_tools_menu)
        
        # Add Convert Documents to JSON option
        document_tools_menu.add_command(
            label="Convert Documents to JSON...", 
            command=self.open_doc_to_json_converter
        )
        
        # Add Split Documents by Headings option
        document_tools_menu.add_command(
            label="Split Documents by Headings...",
            command=self.open_document_splitter
        )
        
        # View menu
        view_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="View", menu=view_menu)
        
        # Add dark mode toggle
        view_menu.add_checkbutton(
            label="Dark Mode", 
            variable=self.dark_mode,
            command=self.update_theme
        )
        
        # Help menu (new addition)
        help_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Help", menu=help_menu)
        
        # Add help dialog option
        help_menu.add_command(
            label="How to Use Block Search",
            command=self.open_help_dialog
        )
        
        # Add keyboard shortcuts option
        help_menu.add_command(
            label="Keyboard Shortcuts",
            command=self.open_shortcuts_help
        )
        
        # Add about option
        help_menu.add_separator()
        help_menu.add_command(
            label="About Block Search",
            command=self.show_about_dialog
        )

    def open_shortcut_settings(self):
        """Open dialog to configure keyboard shortcuts with click-to-record interface."""
        settings_dialog = tk.Toplevel(self.root)
        settings_dialog.title("Keyboard Shortcut Settings")
        settings_dialog.geometry("450x280")
        settings_dialog.resizable(False, False)
        
        # Make dialog modal
        settings_dialog.transient(self.root)
        settings_dialog.grab_set()
        
        # Center dialog on parent
        settings_dialog.update_idletasks()
        screen_width = settings_dialog.winfo_screenwidth()
        screen_height = settings_dialog.winfo_screenheight()
        x = (screen_width - 450) // 2
        y = (screen_height - 280) // 2
        settings_dialog.geometry(f"+{x}+{y}")
        
        # Dialog content
        content_frame = ttk.Frame(settings_dialog, padding=20)
        content_frame.pack(fill=tk.BOTH, expand=True)
        
        # Heading
        heading_label = ttk.Label(content_frame, text="Configure Keyboard Shortcuts", 
                                style="Header.TLabel")
        heading_label.pack(pady=(0, 20))
        
        # Shortcut section
        shortcut_frame = ttk.Frame(content_frame)
        shortcut_frame.pack(fill=tk.X, pady=5)
        
        shortcut_label = ttk.Label(shortcut_frame, text="Focus Search Field:", 
                                 width=18, anchor=tk.W)
        shortcut_label.pack(side=tk.LEFT, padx=(0, 10))
        
        # Current shortcut representation
        current_shortcut = self.format_shortcut_for_display(self.search_focus_shortcut)
        shortcut_var = tk.StringVar(value=current_shortcut)
        
        # Create a style for the recorder in normal and active states
        style = ttk.Style()
        style.configure("Recorder.TFrame", background="#f0f0f0", relief="sunken", borderwidth=1)
        style.configure("Recorder.Active.TFrame", background="#e0e9ff", relief="sunken", borderwidth=2)
        
        # Create a frame to act as the shortcut recorder with a border
        recorder_frame = ttk.Frame(shortcut_frame, style="Recorder.TFrame", width=250, height=30)
        recorder_frame.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 10))
        recorder_frame.pack_propagate(False)  # Fixed size
        
        # Label inside the recorder frame to show the current shortcut
        recorder_label = ttk.Label(recorder_frame, textvariable=shortcut_var,
                                  anchor=tk.CENTER, background="#f0f0f0")
        recorder_label.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5, pady=2)
        
        # Clear button
        clear_button = ttk.Button(shortcut_frame, text="Clear", width=8)
        clear_button.pack(side=tk.LEFT)
        
        # Info label
        info_label = ttk.Label(content_frame, 
                              text="Click in the shortcut field and press your desired key combination.\nPress ESC to cancel recording.",
                              style="Instruction.TLabel")
        info_label.pack(pady=(5, 20))
        
        # Status message
        status_var = tk.StringVar(value="Click in the field to record a shortcut")
        status_label = ttk.Label(content_frame, textvariable=status_var,
                               style="Status.TLabel")
        status_label.pack(pady=(0, 10))
        
        # Buttons frame
        buttons_frame = ttk.Frame(content_frame)
        buttons_frame.pack(fill=tk.X, pady=15)
        
        # Cancel button
        cancel_button = ttk.Button(buttons_frame, text="Cancel", 
                                 command=settings_dialog.destroy)
        cancel_button.pack(side=tk.RIGHT, padx=5)
        
        # Save button
        save_button = ttk.Button(buttons_frame, text="Save", state="disabled")  # Will set command later
        save_button.pack(side=tk.RIGHT, padx=5)
        
        # Shortcut recording variables
        new_shortcut = {"keysym": "", "modifiers": []}
        recording = tk.BooleanVar(value=False)
        
        # Convert the recorded shortcut to Tkinter binding format
        def get_binding_string():
            if not new_shortcut["keysym"]:
                return ""
            
            # Start with the modifiers
            binding = ""
            for mod in sorted(new_shortcut["modifiers"]):
                binding += mod + "-"
                
            # Add the key itself
            binding += new_shortcut["keysym"]
            return binding
        
        # Function to start recording mode
        def start_recording(event=None):
            recording.set(True)
            shortcut_var.set("Press shortcut...")
            recorder_frame.configure(style="Recorder.Active.TFrame")
            recorder_label.configure(background="#e0e9ff")
            status_var.set("Listening for key combination...")
            save_button.config(state="disabled")
            # Clear existing shortcut
            new_shortcut["keysym"] = ""
            new_shortcut["modifiers"] = []
            return "break"  # Prevent default handling
        
        # Function to stop recording mode
        def stop_recording(save_result=False):
            recording.set(False)
            recorder_frame.configure(style="Recorder.TFrame")
            recorder_label.configure(background="#f0f0f0")
            
            if save_result and new_shortcut["keysym"]:
                binding = get_binding_string()
                display = format_shortcut_for_display(binding)
                shortcut_var.set(display)
                save_button.config(state="normal")
                status_var.set("Press Save to apply this shortcut")
            else:
                shortcut_var.set(self.format_shortcut_for_display(self.search_focus_shortcut))
                save_button.config(state="disabled")
                status_var.set("Recording canceled")
        
        # Format a shortcut binding for display
        def format_shortcut_for_display(binding_string):
            if not binding_string:
                return "None"
                
            parts = binding_string.split("-")
            
            # Format for display
            display_parts = []
            
            # Handle modifiers
            for part in parts[:-1]:  # All except the last part
                # Replace Command with ⌘ on macOS
                if part == "Command":
                    display_parts.append("⌘")
                elif part == "Control":
                    display_parts.append("⌃")
                elif part == "Alt":
                    display_parts.append("⌥")
                elif part == "Shift":
                    display_parts.append("⇧")
                else:
                    display_parts.append(part)
            
            # Handle the key itself (last part)
            if parts:
                key = parts[-1]
                # Handle function keys and special keys
                if key.startswith("F") and key[1:].isdigit():
                    display_parts.append(key)  # F1, F2, etc.
                elif key == "space":
                    display_parts.append("Space")  # Space key
                elif len(key) == 1:
                    display_parts.append(key.upper())  # Single letter keys
                else:
                    display_parts.append(key)  # Other keys
            
            return "+".join(display_parts)
        
        # Handler for key press during recording
        def on_key_press(event):
            if not recording.get():
                return
            
            # Get the key itself
            keysym = event.keysym
            
            # Handle special cases
            if keysym == "Escape":
                # Cancel recording
                stop_recording(save_result=False)
                return "break"
            
            # Capture modifiers
            modifiers = []
            if event.state & 0x1:  # Shift
                modifiers.append("Shift")
            if event.state & 0x4:  # Control
                modifiers.append("Control")
            if event.state & 0x8:  # Alt/Option
                modifiers.append("Alt")
            if event.state & 0x38:  # Command (Meta)
                modifiers.append("Command")
            
            # Ignore standalone modifier keys
            if keysym in ["Shift_L", "Shift_R", "Control_L", "Control_R", 
                         "Alt_L", "Alt_R", "Meta_L", "Meta_R"]:
                return "break"
            
            # Special handling for space key
            if keysym == "space":
                keysym = "space"
            
            # Store the new shortcut
            new_shortcut["keysym"] = keysym
            new_shortcut["modifiers"] = modifiers
            
            # End recording with the new shortcut
            stop_recording(save_result=True)
            return "break"  # Prevent default behavior
        
        # Function to clear the shortcut
        def clear_shortcut():
            new_shortcut["keysym"] = ""
            new_shortcut["modifiers"] = []
            shortcut_var.set("None")
            save_button.config(state="normal")
            status_var.set("Shortcut cleared - press Save to apply")
        
        # Save the new shortcut
        def save_shortcut():
            binding = get_binding_string()
            
            # If clearing the shortcut
            if not binding:
                # Unbind old shortcut
                self.root.unbind(self.search_focus_shortcut)
                self.search_focus_shortcut = ""
                status_var.set("Shortcut cleared")
            elif binding:
                # Unbind old shortcut
                self.root.unbind(self.search_focus_shortcut)
                
                # Update and bind new shortcut
                self.search_focus_shortcut = binding
                self.root.bind(self.search_focus_shortcut, self.focus_search_field)  # Use bind consistently
                
                # Show success message
                status_var.set(f"Shortcut saved: {shortcut_var.get()}")
            
            # Save to config
            self.save_config()
            
            # Update menu accelerator
            self.update_shortcut_menu()
            
            save_button.config(state="disabled")
        
        # Configure save button
        save_button.config(command=save_shortcut)
        
        # Configure clear button
        clear_button.config(command=clear_shortcut)
        
        # Make the recorder area clickable to start recording
        recorder_frame.bind("<Button-1>", start_recording)
        recorder_label.bind("<Button-1>", start_recording)
        
        # Bind key events to the dialog
        settings_dialog.bind("<Key>", on_key_press)

    def update_shortcut_menu(self):
        """Update the menu accelerator for the focus search command."""
        if hasattr(self, 'edit_menu') and hasattr(self, 'focus_search_command_index'):
            shortcut_display = self.format_shortcut_for_display(self.search_focus_shortcut)
            self.edit_menu.entryconfigure(
                self.focus_search_command_index, 
                accelerator=shortcut_display if shortcut_display != "None" else ""
            )

    def format_shortcut_for_display(self, binding_string):
        """Convert a Tkinter binding string to a human-readable format."""
        if not binding_string:
            return "None"
            
        parts = binding_string.split("-")
        
        # Format for display
        display_parts = []
        
        # Handle modifiers
        for part in parts[:-1]:  # All except the last part
            # Replace Command with ⌘ on macOS
            if part == "Command":
                display_parts.append("⌘")
            elif part == "Control":
                display_parts.append("⌃")
            elif part == "Alt":
                display_parts.append("⌥")
            elif part == "Shift":
                display_parts.append("⇧")
            else:
                display_parts.append(part)
        
        # Handle the key itself (last part)
        if parts:
            key = parts[-1]
            # Handle function keys and special keys
            if key.startswith("F") and key[1:].isdigit():
                display_parts.append(key)  # F1, F2, etc.
            elif key == "space":
                display_parts.append("Space")  # Space key
            elif len(key) == 1:
                display_parts.append(key.upper())  # Single letter keys
            else:
                display_parts.append(key)  # Other keys
        
        return "+".join(display_parts)

    def configure_styles(self):
        """Configure ttk styles to match macOS appearance with light/dark mode support."""
        style = ttk.Style()
        
        # Choose colors based on mode
        if self.dark_mode.get():
            # Dark mode colors
            bg_color = "#232323"  # Dark background
            accent_color = "#0078FF"  # macOS blue
            text_color = "#FFFFFF"  # White text
            light_text = "#AAAAAA"  # Light gray text
            entry_bg = "#333333"  # Darker for input fields
            tree_bg = "#2A2A2A"  # Slightly lighter for treeview
        else:
            # Light mode colors
            bg_color = "#f5f6f7"  # Light background
            accent_color = "#0077ED"  # macOS blue
            text_color = "#333333"  # Dark text
            light_text = "#737373"  # Gray text
            entry_bg = "#FFFFFF"  # White for input fields
            tree_bg = "#FFFFFF"  # White for treeview
        
        # Configure frame style
        style.configure("Main.TFrame", background=bg_color)
        
        # Configure label styles
        style.configure("Label.TLabel", 
                       background=bg_color, 
                       foreground=text_color, 
                       font=("SF Pro Text", 12))
        
        style.configure("Header.TLabel", 
                       background=bg_color, 
                       foreground=text_color, 
                       font=("SF Pro Display", 18, "bold"))
        
        style.configure("Status.TLabel", 
                       background=bg_color, 
                       foreground=light_text, 
                       font=("SF Pro Text", 11))
        
        style.configure("Instruction.TLabel", 
                       background=bg_color, 
                       foreground=light_text, 
                       font=("SF Pro Text", 11))
        
        # Configure button style
        style.configure("Button.TButton", 
                       font=("SF Pro Text", 12))
        
        # Configure checkbox style
        style.configure("Checkbox.TCheckbutton",
                       background=bg_color,
                       foreground=text_color,
                       font=("SF Pro Text", 12))
        
        # Configure treeview style
        style.configure("Treeview", 
                       background=tree_bg,
                       fieldbackground=tree_bg,
                       foreground=text_color,
                       font=("SF Pro Text", 12), 
                       rowheight=25)
        
        style.configure("Treeview.Heading",
                       background=bg_color,
                       foreground=text_color,
                       font=("SF Pro Text", 12, "bold"))
        
        # Configure treeview selection
        style.map("Treeview", 
                 background=[("selected", accent_color)],
                 foreground=[("selected", "#FFFFFF")])
        
        # Configure entry style
        style.configure("TEntry", 
                       fieldbackground=entry_bg,
                       foreground=text_color)
        
        # Configure combobox style
        style.configure("TCombobox",
                       fieldbackground=entry_bg,
                       background=entry_bg,
                       foreground=text_color,
                       selectbackground=accent_color,
                       selectforeground="#FFFFFF")
        
        # Apply background color to all frames
        for widget in ["TFrame", "TLabelframe"]:
            style.configure(widget, background=bg_color)
    
    def update_theme(self, *args):
        """Update the application theme when dark mode is toggled."""
        # Update styles
        self.configure_styles()
        
        # Update window background
        if self.dark_mode.get():
            # Dark mode
            self.root.configure(background="#232323")
            # Update menu for dark mode appearance
            if sys.platform == 'darwin':  # macOS has special handling for menu bar
                self.root.tk.call('::tk::unsupported::MacWindowStyle', 'appearance', self.root, 'dark')
        else:
            # Light mode
            self.root.configure(background="#f5f6f7")
            # Update menu for light mode appearance
            if sys.platform == 'darwin':  # macOS has special handling for menu bar
                self.root.tk.call('::tk::unsupported::MacWindowStyle', 'appearance', self.root, 'aqua')
        
        # Force refresh of Treeview styles
        self.tree.tag_configure("refresh", background="")
        self.tree.update_idletasks() # Ensures all UI updates are applied

        # Save the preference
        self.save_config()
        
        # Update status
        mode_name = "Dark Mode" if self.dark_mode.get() else "Light Mode"
        self.status_var.set(f"Switched to {mode_name}")

    def open_help_dialog(self):
        """Open the comprehensive help dialog."""
        HelpDialog(self.root)

    def open_shortcuts_help(self):
        """Open the help dialog directly to the shortcuts tab."""
        help_dialog = HelpDialog(self.root)
        help_dialog.notebook.select(4)  # Index of the shortcuts tab

    def show_about_dialog(self):
        """Show information about the application."""
        about_text = (
            "Block Search\n\n"
            "Version 1.0\n\n"
            "A powerful tool for document management, conversion, and content searching.\n\n"
            "Features:\n"
            "• Convert between JSON and Word formats\n"
            "• Search and filter document content\n"
            "• Split documents by headings\n"
            "• Batch document processing\n\n"
            "© 2025 Your Organization"
        )
        
        # Get logo image if available
        try:
            logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "resources", "logo.png")
            if os.path.exists(logo_path):
                from PIL import Image, ImageTk
                logo = Image.open(logo_path)
                logo = logo.resize((100, 100), Image.LANCZOS)
                logo_image = ImageTk.PhotoImage(logo)
            else:
                logo_image = None
        except ImportError:
            logo_image = None
        
        # Create about dialog
        about_dialog = tk.Toplevel(self.root)
        about_dialog.title("About Block Search")
        about_dialog.geometry("400x450")
        about_dialog.resizable(False, False)
        
        # Make dialog modal
        about_dialog.transient(self.root)
        about_dialog.grab_set()
        
        # Center on parent
        about_dialog.update_idletasks()
        parent_x = self.root.winfo_x()
        parent_y = self.root.winfo_y()
        parent_width = self.root.winfo_width()
        parent_height = self.root.winfo_height()
        
        width = about_dialog.winfo_width()
        height = about_dialog.winfo_height()
        
        x = parent_x + (parent_width - width) // 2
        y = parent_y + (parent_height - height) // 2
        
        about_dialog.geometry(f"+{x}+{y}")
        
        # Apply consistent color scheme
        bg_color = "#F0F0F0"
        fg_color = "#000000"
        
        about_dialog.configure(bg=bg_color)
        
        # Add logo if available
        if logo_image:
            logo_label = tk.Label(about_dialog, image=logo_image, bg=bg_color)
            logo_label.image = logo_image  # Keep a reference to prevent garbage collection
            logo_label.pack(pady=(20, 10))
        
        # Add about text
        about_label = tk.Label(
            about_dialog,
            text=about_text,
            font=("Helvetica", 12),
            bg=bg_color,
            fg=fg_color,
            justify=tk.CENTER,
            padx=20,
            pady=20
        )
        about_label.pack(expand=True)
        
        # Add close button
        close_button = tk.Button(
            about_dialog,
            text="Close",
            command=about_dialog.destroy,
            bg="#E0E0E0",
            fg="#000000",
            font=("Helvetica", 12),
            width=10
        )
        close_button.pack(pady=20)

    def open_document_splitter(self):
        """Open the dialog for splitting Word documents by headings."""
        DocSplitterDialog(self.root)

    def open_doc_to_json_converter(self):
        """Open the dialog for batch converting Word documents to JSON."""
        DocToJSONConverter(self.root, self.dark_mode.get())

    def load_config(self):
        """Load settings from configuration file."""
        # Create config directory if it doesn't exist
        config_dir = os.path.dirname(self.CONFIG_FILE)
        if not os.path.exists(config_dir):
            os.makedirs(config_dir, exist_ok=True)
        
        # Load existing config if available
        if os.path.exists(self.CONFIG_FILE):
            try:
                with open(self.CONFIG_FILE, "r") as f:
                    config = json.load(f)
                    
                # Set directory if it exists in config
                if "directory" in config and os.path.exists(config["directory"]):
                    self.directory_var.set(config["directory"])
                
                # Set template if it exists in config
                if "template" in config and os.path.exists(config["template"]):
                    self.template_var.set(config["template"])
                    
                # Set target document if it exists in config
                if "target_document" in config:
                    saved_target = config["target_document"]
                    # Will be applied after refreshing document list
                    
                # Set dark mode preference if it exists
                if "dark_mode" in config:
                    self.dark_mode.set(config["dark_mode"])
                    
            except Exception as e:
                print(f"Error loading config: {e}")
        
        # Check for default template in directory
        if self.directory_var.get() and not self.template_var.get():
            default_template = os.path.join(self.directory_var.get(), self.DEFAULT_TEMPLATE_NAME)
            if os.path.exists(default_template):
                self.template_var.set(default_template)
        
        # Refresh document list and try to select saved document
        self.refresh_word_documents()
        
        # Try to select saved document if it exists
        if 'saved_target' in locals() and saved_target in self.open_documents:
            self.target_document_var.set(saved_target)
        
        # Initialize file list if directory is set
        if self.directory_var.get():
            self.search_files()

    def save_config(self):
        """Save settings to configuration file."""
        config = {
            "directory": self.directory_var.get(),
            "template": self.template_var.get(),
            "target_document": self.target_document_var.get(),
            "dark_mode": self.dark_mode.get()
        }
        
        try:
            with open(self.CONFIG_FILE, "w") as f:
                json.dump(config, f)
        except Exception as e:
            print(f"Error saving config: {e}")
    
    def handle_search_enter_key(self, event):
        """Handle enter key in search field - trigger selection of highlighted item."""
        # Get the current selection
        selection = self.tree.selection()
        if selection:
            # If any item is selected, trigger the on_file_select method
            self.on_file_select(event)
        return "break"  # Prevent default behavior

    def handle_search_down_key(self, event):
        """Handle down arrow key in search field - move to first item."""
        items = self.tree.get_children()
        if items:
            # Select first item
            self.tree.selection_set(items[1])
            self.tree.focus_set()
            self.tree.focus(items[1])
            # Ensure it's visible
            self.tree.see(items[1])
        return "break"  # Prevent default behavior

    def handle_search_up_key(self, event):
        """Handle up arrow key in search field - move to last item."""
        items = self.tree.get_children()
        if items:
            # Select last item
            last_item = items[-1]
            self.tree.selection_set(last_item)
            self.tree.focus_set()
            self.tree.focus(last_item)
            # Ensure it's visible
            self.tree.see(last_item)
        return "break"  # Prevent default behavior

    def handle_tree_navigation(self, event, direction):
        """Handle up/down navigation in the tree."""
        # Get current selection
        selection = self.tree.selection()
        if not selection:
            # If nothing is selected, select the first or last item
            items = self.tree.get_children()
            if not items:
                return "break"
                
            if direction == "up":
                item_to_select = items[-1]  # Last item
            else:
                item_to_select = items[0]  # First item
        else:
            # Get current selected item
            current_item = selection[0]
            current_index = self.tree.index(current_item)
            items = self.tree.get_children()
            
            # Determine next item based on direction
            if direction == "up":
                if current_index > 0:
                    item_to_select = items[current_index - 1]
                else:
                    item_to_select = items[-1]  # Wrap to end
            else:  # direction == "down"
                if current_index < len(items) - 1:
                    item_to_select = items[current_index + 1]
                else:
                    item_to_select = items[0]  # Wrap to beginning
        
        # Select the item
        self.tree.selection_set(item_to_select)
        self.tree.focus(item_to_select)
        self.tree.see(item_to_select)
        return "break"  # Prevent default behavior

    def focus_search_field(self, event=None):
        """Move focus to search field and select all text."""
        if hasattr(self, 'search_entry') and self.search_entry.winfo_exists():
            # Focus on search entry
            self.search_entry.focus_set()
            
            # Select all text for easy replacement
            self.search_entry.selection_range(0, tk.END)
            
        return "break"  # Prevent default behavior

    def refresh_word_documents(self, event=None):
        """Refresh the list of open Word documents using AppleScript."""
        self.status_var.set("Refreshing Word documents...")
        self.root.update_idletasks()
        
        try:
            # Simplified AppleScript that has better compatibility
            applescript = '''
            tell application "Microsoft Word"
                set docNames to {}
                if (count of documents) > 0 then
                    repeat with i from 1 to (count of documents)
                        set end of docNames to name of document i
                    end repeat
                end if
                return docNames
            end tell
            '''
            
            # Run the AppleScript
            process = subprocess.Popen(['osascript', '-e', applescript], 
                                      stdout=subprocess.PIPE, 
                                      stderr=subprocess.PIPE)
            stdout, stderr = process.communicate()
            
            if process.returncode != 0:
                error_msg = stderr.decode().strip()
                print(f"AppleScript error: {error_msg}")  # Print for debugging
                messagebox.showerror("Error", f"Failed to get Word documents:\n{error_msg}")
                return
            
            # Parse the output - AppleScript returns comma-separated items
            documents_output = stdout.decode().strip()
            
            # Default option always available
            document_list = ["Copy to clipboard only"]
            
            # Add open documents if there are any
            if documents_output:
                # Remove any curly braces that might be in the output
                documents_output = documents_output.replace('{', '').replace('}', '')
                docs = [doc.strip() for doc in documents_output.split(',')]
                for doc in docs:
                    if doc and doc != '':
                        document_list.append(doc)
            
            # Update the dropdown
            self.open_documents = document_list
            self.target_dropdown['values'] = document_list
            
            # Keep the same selection if possible
            current_selection = self.target_document_var.get()
            if current_selection not in document_list:
                self.target_document_var.set("Copy to clipboard only")
                
            self.status_var.set(f"Found {len(document_list) - 1} open Word documents")
            
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            print(f"Python error: {error_details}")  # Print for debugging
            messagebox.showerror("Error", str(e))
            self.status_var.set("Error refreshing documents")

    def browse_directory(self):
        """Open file dialog to select directory."""
        directory = filedialog.askdirectory()
        if directory:
            self.directory_var.set(directory)
            
            # Check for default template in the directory
            default_template = os.path.join(directory, self.DEFAULT_TEMPLATE_NAME)
            if os.path.exists(default_template) and not self.template_var.get():
                self.template_var.set(default_template)
                
            # Update file list and save config
            self.search_files()
            self.save_config()
    
    def browse_template(self):
        """Open file dialog to select template document."""
        template_file = filedialog.askopenfilename(
            title="Select Template Document",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        
        if template_file:
            self.template_var.set(template_file)
            self.save_config()
    
    def search_files(self, *args):
        """
        Search for JSON files in the selected directory with enhanced search capabilities.
        - Splits search term into words
        - Matches each word individually (partial matching enabled)
        - All words must be present in the filename to be a match (AND logic)
        - Case insensitive
        - Selects the first result automatically
        """
        directory = self.directory_var.get()
        search_string = self.search_var.get().strip()
        
        # Clear the treeview
        for item in self.tree.get_children():
            self.tree.delete(item)
        
        if not directory or not os.path.isdir(directory):
            return
        
        self.status_var.set("Searching...")
        self.root.update_idletasks()
        
        # Split search string into individual words
        search_terms = [term.lower() for term in search_string.split() if term]
        
        file_count = 0
        file_data = []
        
        # Walk through directory tree
        for root, _, files in os.walk(directory):
            for filename in files:
                if not filename.lower().endswith('.json'):
                    continue
                    
                # Check if all search terms are in the filename
                filename_lower = filename.lower()
                
                # If there are no search terms, include all files
                if not search_terms:
                    is_match = True
                else:
                    # Check if all search terms are found in the filename
                    is_match = all(term in filename_lower for term in search_terms)
                
                if is_match:
                    file_count += 1
                    file_path = os.path.join(root, filename)
                    
                    # Get file last accessed time (last opened)
                    file_stat = os.stat(file_path)
                    access_time = file_stat.st_atime
                    access_time_str = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(access_time))
                    
                    # Calculate relative path for display
                    rel_path = os.path.relpath(root, directory)
                    if rel_path == ".":
                        display_name = filename[:-5]  # Remove .json extension
                        rel_path_display = "<root>"
                    else:
                        display_name = filename[:-5]
                        rel_path_display = rel_path
                    
                    # Store all data in a list for later sorting
                    file_data.append({
                        "filename": display_name,
                        "relpath": rel_path_display,
                        "lastopened": access_time_str,
                        "access_timestamp": access_time,
                        "full_path": file_path
                    })
        
        # Apply current sort
        self.sort_file_data(file_data)
        
        # Add sorted items to treeview
        for item in file_data:
            self.tree.insert("", tk.END, values=(
                item["filename"], 
                item["relpath"], 
                item["lastopened"]
            ))
        
        # Select the first item if there are any results
        items = self.tree.get_children()
        if items:
            first_item = items[0]
            self.tree.selection_set(first_item)
            # Note: We don't change focus here to allow continuous typing in search field
        
        self.status_var.set(f"Found {file_count} matching files")
    
    def sort_file_data(self, file_data):
        """Sort the file data based on current sort settings."""
        column = self.current_sort["column"]
        reverse = self.current_sort["reverse"] or self.reverse_sort_var.get()
        
        if column == "lastopened":
            # Sort by timestamp for last opened date
            file_data.sort(key=lambda x: x["access_timestamp"], reverse=reverse)
        else:
            # Sort by string for other columns
            file_data.sort(key=lambda x: x[column].lower(), reverse=reverse)

    def sort_treeview(self, column, refresh=True):
        """Sort treeview data by the specified column."""
        # Update the current sort
        if self.current_sort["column"] == column and refresh is False:
            # Toggle reverse if clicking the same column again
            self.current_sort["reverse"] = not self.current_sort["reverse"]
        else:
            self.current_sort["column"] = column
            if refresh is False:
                self.current_sort["reverse"] = False
        
        # Perform the search again to reload and sort items
        self.search_files()
        
        # Update column headings to show sort direction
        for col in ("filename", "relpath", "lastopened"):
            if col == column:
                direction = " ↑" if not (self.current_sort["reverse"] or self.reverse_sort_var.get()) else " ↓"
                # Format column header text
                if col == "filename":
                    header_text = "File Name"
                elif col == "relpath":
                    header_text = "Relative Path"
                elif col == "lastopened":
                    header_text = "Last Opened"
                else:
                    header_text = col.capitalize()
                    
                self.tree.heading(col, text=f"{header_text}{direction}", 
                                 command=lambda c=col: self.sort_treeview(c, False))
            else:
                # Format column header text
                if col == "filename":
                    header_text = "File Name"
                elif col == "relpath":
                    header_text = "Relative Path"
                elif col == "lastopened":
                    header_text = "Last Opened"
                else:
                    header_text = col.capitalize()
                    
                self.tree.heading(col, text=header_text, 
                                 command=lambda c=col: self.sort_treeview(c, False))

    def refresh_sort(self):
        """Refresh the current sort when reverse checkbox is toggled."""
        self.sort_treeview(self.current_sort["column"], True)

    def convert_json_to_docx(self, json_file, output_docx):
        """Convert JSON to DOCX using Pandoc, using a proper temporary directory."""
        # Create a temporary directory for the conversion process
        with tempfile.TemporaryDirectory() as temp_dir:
            # Use the temporary directory for the output file
            temp_output = os.path.join(temp_dir, "temp_output.docx")
            
            pandoc_path = get_pandoc_path()
            print(f"Using pandoc at: {pandoc_path}")
            
            command = [
                pandoc_path,
                "-f", "json",
                "-t", "docx",
                "--wrap=none",
                "--extract-media=.",
                "--preserve-tabs",
                "-o", temp_output,
                json_file
            ]
            
            # Add template reference if available
            template_path = self.template_var.get()
            if template_path and os.path.exists(template_path):
                command.insert(5, f"--reference-doc={template_path}")
            
            try:
                self.status_var.set("Converting JSON to DOCX...")
                self.root.update_idletasks()
                
                print(f"Running command: {' '.join(command)}")
                result = subprocess.run(command, capture_output=True, text=True)
                
                if result.returncode != 0:
                    error_msg = result.stderr
                    print(f"Pandoc error: {error_msg}")
                    messagebox.showerror("Conversion Error", 
                                       f"Pandoc conversion failed:\n{error_msg}")
                    return False
                
                # Copy the temp file to the final location if needed
                if temp_output != output_docx:
                    import shutil
                    shutil.copy2(temp_output, output_docx)
                
                return True
            except Exception as e:
                import traceback
                error_details = traceback.format_exc()
                print(f"Exception during conversion: {error_details}")
                messagebox.showerror("Conversion Error", str(e))
                return False
    
    def copy_docx_to_clipboard_using_applescript(self, doc_path):
        """
        Copy document content to clipboard using Word's native AppleScript commands.
        Uses a simplified approach that works with the active document.
        """
        print(f"\n*** Starting clipboard copy operation for: {doc_path}")
        
        if not os.path.exists(doc_path):
            messagebox.showerror("Error", f"Word file not found:\n{doc_path}")
            return False
        
        # Get absolute path for AppleScript
        abs_path = os.path.abspath(doc_path)
        
        self.status_var.set("Opening document in Word and copying content...")
        self.root.update_idletasks()
        
        # Simplified AppleScript that doesn't rely on document references
        applescript = f'''
        tell application "Microsoft Word"
            activate
            
            -- Open the document
            open POSIX file "{abs_path}"
            delay 1  -- Allow document to fully load

            -- Ensure the document is active
            set theDoc to active document

            -- Select all content properly
            set myRange to text object of theDoc
            select myRange
            delay 0.5  -- Small delay to ensure selection

            -- Copy selection to clipboard
            copy object selection
            delay 1  -- Ensure clipboard updates before closing
            
            -- Close document without saving
            close theDoc saving no
        end tell
        '''
        
        try:
            print("*** Running AppleScript to copy content")
            # Run the AppleScript
            process = subprocess.Popen(['osascript', '-e', applescript], 
                                      stdout=subprocess.PIPE, 
                                      stderr=subprocess.PIPE)
            stdout, stderr = process.communicate()
            
            if process.returncode != 0:
                error_msg = stderr.decode().strip()
                print(f"*** AppleScript ERROR: {error_msg}")
                messagebox.showerror("Error", 
                                   f"Failed to copy content from Word:\n{stderr.decode()}")
                return False
            
            # Verify clipboard has content
            time.sleep(0.5) 
            
            self.status_var.set("Content copied to clipboard successfully")
            print("*** Copy operation completed successfully")
            return True
                
        except Exception as e:
            print(f"*** Exception during copy operation: {e}")
            messagebox.showerror("Error", str(e))
            self.status_var.set("Error occurred")
            return False

    def paste_to_specific_document(self, document_name, doc_path):
        """Activate a specific Word document and paste content using System Events keystrokes."""
        print(f"\n+++ Starting paste operation to document: {document_name}")
        
        if document_name == "Copy to clipboard only":
            return True
                    
        try:
            # Original method using System Events to send Cmd+V
            applescript = f'''
            tell application "Microsoft Word"
                set pasted to false
                
                -- Find and activate the document by name
                repeat with i from 1 to count of documents
                    if name of document i is "{document_name}" then
                        activate
                        set active document to document i
                        delay 0.5 -- Allow Word to focus properly
                        
                        -- Paste using Word's built-in paste command with formatting option
                        paste and format (text object of selection) type paste default
                        
                        set pasted to true
                        exit repeat
                    end if
                end repeat
                
                return pasted
            end tell
            '''
            
            print("+++ Running AppleScript to paste content")
            # Run the AppleScript
            process = subprocess.Popen(['osascript', '-e', applescript], 
                                      stdout=subprocess.PIPE, 
                                      stderr=subprocess.PIPE)
            stdout, stderr = process.communicate()
            
            # Check result
            if process.returncode != 0:
                error_msg = stderr.decode().strip()
                print(f"+++ AppleScript ERROR: {error_msg}")
                
                # Specifically handle the permissions error
                if "not allowed to send keystrokes" in error_msg or "1002" in str(error_msg):
                    # This is a permissions issue - prompt the user
                    response = messagebox.askquestion(
                        "Permission Required",
                        "BlockSearch needs permission to paste content automatically.\n\n"
                        "Do you want to open Accessibility settings to grant this permission?\n\n"
                        "(Your content is already copied to clipboard and can be pasted manually)",
                        icon='info'
                    )
                    
                    if response == 'yes':
                        # Open the Security & Privacy preferences directly to Accessibility
                        open_script = '''
                        tell application "System Preferences"
                            activate
                            set current pane to pane id "com.apple.preference.security"
                            delay 0.5
                            tell application "System Events" to tell process "System Preferences"
                                delay 0.5
                                if exists tab group 1 of window "Security & Privacy" then
                                    click radio button "Privacy" of tab group 1 of window "Security & Privacy"
                                    delay 0.3
                                    if exists row "Accessibility" of table 1 of scroll area 1 of tab group 1 of window "Security & Privacy" then
                                        select row "Accessibility" of table 1 of scroll area 1 of tab group 1 of window "Security & Privacy"
                                    end if
                                end if
                            end tell
                        end tell
                        '''
                        
                        try:
                            subprocess.Popen(['osascript', '-e', open_script], 
                                            stdout=subprocess.PIPE, 
                                            stderr=subprocess.PIPE)
                        except Exception as e:
                            print(f"Error opening System Preferences: {e}")
                    
                    # Activate Word so the user can paste manually
                    activate_script = f'''
                    tell application "Microsoft Word"
                        activate
                        repeat with i from 1 to (count of documents)
                            if name of document i is "{document_name}" then
                                set active document to document i
                                exit repeat
                            end if
                        end repeat
                    end tell
                    '''
                    
                    try:
                        subprocess.Popen(['osascript', '-e', activate_script], 
                                        stdout=subprocess.PIPE, 
                                        stderr=subprocess.PIPE)
                        self.status_var.set(f"Permission needed. Please paste manually with Cmd+V")
                    except Exception as e:
                        print(f"Error activating Word: {e}")
                        
                return False
                    
            # Check if paste was successful
            result = stdout.decode().strip()
            success = result.lower() == "true"
            print(f"+++ Paste operation result: {success}")
            
            # Clean up temporary file
            if os.path.exists(doc_path):
                print(f"*** Removing temporary file: {doc_path}")
                try:
                    os.remove(doc_path)
                except Exception as cleanup_error:
                    print(f"*** Warning: Could not remove temp file: {cleanup_error}")
            
            return success

        except Exception as e:
            print(f"+++ Exception during paste operation: {e}")
            return False

    def on_file_select(self, event):
        """Handle file selection from treeview with proper temp file handling."""
        # Debug info
        print(f"\n>>> on_file_select called with event type: {type(event).__name__}")

        # Activate Busy Flag
        if hasattr(self, '_busy') and self._busy:
            print("!!! Already busy, ignoring duplicate call")
            return
        self._busy = True
        print(">>> Setting busy flag, starting processing")
        
        try:
            selected_items = self.tree.selection()
            if not selected_items:
                print(">>> No items selected, returning")
                return
            
            selected_item = selected_items[0]
            values = self.tree.item(selected_item, "values")
            display_name = values[0]  # Filename
            rel_path = values[1]      # Relative path
            
            print(f">>> Selected: {display_name} in {rel_path}")
            
            # Prevent duplicate execution within a short time
            current_time = time.time()
            if hasattr(self, "_last_click_time") and (current_time - self._last_click_time < 15):
                if hasattr(self, "_last_clicked_item") and (self._last_clicked_item == display_name):
                    print(">>> ignoring duplicate event due to debounce protection")
                    return
            self._last_click_time = current_time  # Update last click time
            self._last_clicked_item = display_name # Update last clicked name

            # Construct full path to JSON file
            directory = self.directory_var.get()
            json_filename = display_name + ".json"
            
            if rel_path == "<root>":
                json_path = os.path.join(directory, json_filename)
            else:
                json_path = os.path.join(directory, rel_path, json_filename)
            
            if not os.path.exists(json_path):
                messagebox.showerror("Error", f"File not found:\n{json_path}")
                return
            
            # Create temp file in a proper temporary directory
            temp_dir = tempfile.gettempdir()
            output_docx = os.path.join(temp_dir, "temp_output.docx")
            
            # Convert and copy to clipboard
            self.status_var.set(f"Processing {json_filename}...")
            
            print(f">>> Converting JSON to DOCX: {json_path} -> {output_docx}")
            if self.convert_json_to_docx(json_path, output_docx):
                print(f">>> Conversion successful, copying to clipboard")
                # Use the AppleScript-based clipboard method 
                if self.copy_docx_to_clipboard_using_applescript(output_docx):
                    # If target is not clipboard-only, try to paste to Word
                    target_doc = self.target_document_var.get()
                    if target_doc != "Copy to clipboard only":
                        print(f">>> Attempting to paste to document: {target_doc}")
                        self.status_var.set(f"Pasting to document: {target_doc}...")
                        self.root.update_idletasks()
                        
                        # Use the combined function
                        if self.paste_to_specific_document(target_doc, output_docx):
                            self.status_var.set(f"Content pasted to '{target_doc}'")
                        else:
                            self.status_var.set("Failed to paste, but content is on clipboard")
                            messagebox.showinfo(
                                "Paste Failed", 
                                f"Could not paste to '{target_doc}'.\n\n"
                                f"The content has been copied to clipboard. You can manually paste it."
                            )
                    else:
                        self.status_var.set("Content copied to clipboard successfully")
                else:
                    self.status_var.set("Failed to copy to clipboard")
                    
                # Clean up the temporary file
                try:
                    if os.path.exists(output_docx):
                        os.remove(output_docx)
                except Exception as e:
                    print(f">>> Warning: Failed to clean up temp file: {e}")
        finally:
            print(">>> Clearing busy flag")
            self._busy = False

    def on_double_click(self, event):
        """Catch double clicks and pass them on to on_file_select with debug message."""
        print(f"\n>>> on_double_click called with event type: {type(event).__name__}")

        self.on_file_select(event)

def check_pandoc():
    """Check if Pandoc is installed or available in the app bundle."""
    # If we're in a bundled app, try to find the bundled pandoc
    if getattr(sys, 'frozen', False):
        pandoc_path = get_pandoc_path()
        if pandoc_path != "pandoc":  # If we found a specific path
            print(f"Using bundled pandoc at: {pandoc_path}")
            return True
    
    # Otherwise check if pandoc is in PATH
    try:
        result = subprocess.run(["pandoc", "--version"], 
                              stdout=subprocess.PIPE, 
                              stderr=subprocess.PIPE)
        return result.returncode == 0
    except FileNotFoundError:
        return False

class DocToJSONConverter(tk.Toplevel):
    """Dialog for batch converting Word documents to JSON format with style preservation."""
    
    def __init__(self, parent, dark_mode=False):
        super().__init__(parent)
        self.parent = parent
        
        # Set dialog properties
        self.title("Convert Documents to JSON")
        self.geometry("650x450")
        self.resizable(True, True)
        self.minsize(600, 400)
        
        # Make dialog modal
        self.transient(parent)
        self.grab_set()
        
        # Center on parent
        self.update_idletasks()
        parent_x = parent.winfo_x()
        parent_y = parent.winfo_y()
        parent_width = parent.winfo_width()
        parent_height = parent.winfo_height()
        
        width = self.winfo_width()
        height = self.winfo_height()
        
        x = parent_x + (parent_width - width) // 2
        y = parent_y + (parent_height - height) // 2
        
        self.geometry(f"+{x}+{y}")
        
        # Variables
        self.input_folder_var = tk.StringVar()
        self.output_folder_var = tk.StringVar()
        self.template_var = tk.StringVar()
        self.status_var = tk.StringVar(value="Ready")
        self.progress_var = tk.DoubleVar(value=0.0)
        self.is_running = False
        self.conversion_thread = None
        
        # Apply simple high-contrast colors
        self.configure(background="#F0F0F0")  # Light gray background
        
        # Setup UI
        self.setup_ui()
        
        # Setup dialog close handling
        self.protocol("WM_DELETE_WINDOW", self.on_close)
        
    def setup_ui(self):
        """Create the user interface with simplified styling."""
        # Create main frame with padding
        main_frame = tk.Frame(self, bg="#F0F0F0", padx=20, pady=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Description text at the top
        description_text = (
            "Convert Word documents (.docx) to JSON format with style preservation. "
            "This tool processes all .docx files in the input folder and saves JSON outputs to the specified location.\n"
            "TIP: For best style preservation, use a completely blank Verbatim document as the template."
        )
        
        description_label = tk.Label(
            main_frame, 
            text=description_text,
            bg="#F0F0F0",
            fg="#000000",
            font=("Helvetica", 12),
            wraplength=600,
            justify=tk.LEFT
        )
        description_label.pack(fill=tk.X, pady=(0, 20))
        
        # Input folder selection
        input_frame = tk.Frame(main_frame, bg="#F0F0F0")
        input_frame.pack(fill=tk.X, pady=5)
        
        input_label = tk.Label(
            input_frame, 
            text="Input Folder:", 
            bg="#F0F0F0",
            fg="#000000",
            font=("Helvetica", 12),
            width=15,
            anchor=tk.W
        )
        input_label.pack(side=tk.LEFT)
        
        input_entry = tk.Entry(
            input_frame, 
            textvariable=self.input_folder_var,
            bg="#FFFFFF",
            fg="#000000",
            font=("Helvetica", 12)
        )
        input_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(5, 5))
        
        input_button = tk.Button(
            input_frame, 
            text="Browse...", 
            command=self.browse_input_folder,
            bg="#E0E0E0",
            fg="#000000",
            font=("Helvetica", 12),
            relief=tk.RAISED
        )
        input_button.pack(side=tk.LEFT)
        
        # Output folder selection
        output_frame = tk.Frame(main_frame, bg="#F0F0F0")
        output_frame.pack(fill=tk.X, pady=5)
        
        output_label = tk.Label(
            output_frame, 
            text="Output Folder:", 
            bg="#F0F0F0",
            fg="#000000",
            font=("Helvetica", 12),
            width=15,
            anchor=tk.W
        )
        output_label.pack(side=tk.LEFT)
        
        output_entry = tk.Entry(
            output_frame, 
            textvariable=self.output_folder_var,
            bg="#FFFFFF",
            fg="#000000",
            font=("Helvetica", 12)
        )
        output_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(5, 5))
        
        output_button = tk.Button(
            output_frame, 
            text="Browse...", 
            command=self.browse_output_folder,
            bg="#E0E0E0",
            fg="#000000",
            font=("Helvetica", 12),
            relief=tk.RAISED
        )
        output_button.pack(side=tk.LEFT)
        
        # Template selection
        template_frame = tk.Frame(main_frame, bg="#F0F0F0")
        template_frame.pack(fill=tk.X, pady=5)
        
        template_label = tk.Label(
            template_frame, 
            text="Template Doc:", 
            bg="#F0F0F0",
            fg="#000000",
            font=("Helvetica", 12),
            width=15,
            anchor=tk.W
        )
        template_label.pack(side=tk.LEFT)
        
        template_entry = tk.Entry(
            template_frame, 
            textvariable=self.template_var,
            bg="#FFFFFF",
            fg="#000000",
            font=("Helvetica", 12)
        )
        template_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(5, 5))
        
        template_button = tk.Button(
            template_frame, 
            text="Browse...", 
            command=self.browse_template,
            bg="#E0E0E0",
            fg="#000000",
            font=("Helvetica", 12),
            relief=tk.RAISED
        )
        template_button.pack(side=tk.LEFT)
        
        # Progress section
        progress_frame = tk.Frame(main_frame, bg="#F0F0F0")
        progress_frame.pack(fill=tk.X, pady=(20, 5))
        
        # Use a standard progressbar without styling
        self.progress_bar = ttk.Progressbar(
            progress_frame,
            orient="horizontal",
            mode="determinate",
            variable=self.progress_var
        )
        self.progress_bar.pack(fill=tk.X, pady=5)
        
        status_label = tk.Label(
            progress_frame, 
            textvariable=self.status_var,
            bg="#F0F0F0",
            fg="#505050",
            font=("Helvetica", 11)
        )
        status_label.pack(anchor=tk.W)
        
        # Buttons frame at bottom
        buttons_frame = tk.Frame(main_frame, bg="#F0F0F0")
        buttons_frame.pack(fill=tk.X, pady=(20, 0), side=tk.BOTTOM)
        
        self.convert_button = tk.Button(
            buttons_frame, 
            text="Convert", 
            command=self.start_conversion,
            bg="#E0E0E0",
            fg="#000000",
            font=("Helvetica", 12),
            relief=tk.RAISED,
            padx=10
        )
        self.convert_button.pack(side=tk.RIGHT, padx=5)
        
        self.cancel_button = tk.Button(
            buttons_frame, 
            text="Cancel", 
            command=self.on_close,
            bg="#E0E0E0",
            fg="#000000",
            font=("Helvetica", 12),
            relief=tk.RAISED,
            padx=10
        )
        self.cancel_button.pack(side=tk.RIGHT, padx=5)
    
    def browse_input_folder(self):
        """Open directory dialog to select input folder."""
        folder = filedialog.askdirectory(title="Select Input Folder")
        if folder:
            self.input_folder_var.set(folder)
    
    def browse_output_folder(self):
        """Open directory dialog to select output folder."""
        folder = filedialog.askdirectory(title="Select Output Folder")
        if folder:
            self.output_folder_var.set(folder)
    
    def browse_template(self):
        """Open file dialog to select template document."""
        template_file = filedialog.askopenfilename(
            title="Select Template Document",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        if template_file:
            self.template_var.set(template_file)
    
    def validate_inputs(self):
        """Validate that all required inputs are provided."""
        if not self.input_folder_var.get():
            messagebox.showerror("Input Error", "Please select an input folder.")
            return False
        
        if not os.path.isdir(self.input_folder_var.get()):
            messagebox.showerror("Input Error", "Input folder does not exist.")
            return False
        
        if not self.output_folder_var.get():
            messagebox.showerror("Input Error", "Please select an output folder.")
            return False
        
        # Create output folder if it doesn't exist
        if not os.path.exists(self.output_folder_var.get()):
            try:
                os.makedirs(self.output_folder_var.get())
            except OSError as e:
                messagebox.showerror("Output Error", f"Could not create output folder: {str(e)}")
                return False
        
        # Check if template exists if specified
        if self.template_var.get() and not os.path.isfile(self.template_var.get()):
            messagebox.showerror("Template Error", "Template file does not exist.")
            return False
        
        return True
    
    def start_conversion(self):
        """Start the conversion process in a separate thread."""
        if self.is_running:
            return
        
        if not self.validate_inputs():
            return
        
        # Update UI
        self.convert_button.config(state="disabled")
        self.status_var.set("Starting conversion...")
        self.progress_var.set(0)
        self.is_running = True
        
        # Start conversion in a separate thread
        import threading
        self.conversion_thread = threading.Thread(target=self.convert_documents)
        self.conversion_thread.daemon = True
        self.conversion_thread.start()
    
    def convert_documents(self):
        """Convert all docx files in the input folder to JSON."""
        try:
            input_folder = self.input_folder_var.get()
            output_folder = self.output_folder_var.get()
            template = self.template_var.get()
            
            # Get list of docx files
            docx_files = []
            for file in os.listdir(input_folder):
                if file.lower().endswith('.docx'):
                    docx_files.append(os.path.join(input_folder, file))
            
            if not docx_files:
                self.update_status("No .docx files found in the input folder.", True)
                return
            
            # Get the pandoc path from our helper function
            pandoc_path = get_pandoc_path()
            print(f"DocToJSONConverter: Using pandoc at: {pandoc_path}")
            
            # Verify pandoc executable exists and is executable
            if not os.path.isfile(pandoc_path) and pandoc_path != "pandoc":
                error_msg = f"Pandoc executable not found at: {pandoc_path}"
                print(f"ERROR: {error_msg}")
                self.update_status(f"Error: {error_msg}", True)
                return
                
            if pandoc_path != "pandoc" and not os.access(pandoc_path, os.X_OK):
                error_msg = f"Pandoc exists but is not executable: {pandoc_path}"
                print(f"ERROR: {error_msg}")
                # Try to fix permissions
                try:
                    print("Attempting to fix permissions...")
                    os.chmod(pandoc_path, 0o755)  # rwxr-xr-x
                    print(f"New permissions set on {pandoc_path}")
                except Exception as perm_error:
                    print(f"Failed to set permissions: {perm_error}")
                    self.update_status(f"Error: {error_msg}", True)
                    return
            
            # Setup conversion command base
            base_command = [
                pandoc_path,  # Use the proper pandoc path
                "-f", "docx+styles",    # Use styles extension to preserve custom styles
                "-t", "json",
                "--wrap=none"           # Preserve wrapping which can affect formatting
            ]
            
            # Add template if specified
            if template:
                base_command.extend(["--reference-doc", template])
            
            # Process each file
            total_files = len(docx_files)
            for i, docx_file in enumerate(docx_files):
                # Skip template file if it's in the input directory
                if os.path.abspath(docx_file) == os.path.abspath(template):
                    continue
                
                # Update status
                file_name = os.path.basename(docx_file)
                self.update_status(f"Converting {i+1}/{total_files}: {file_name}")
                
                # Create output file path
                output_file = os.path.join(
                    output_folder, 
                    os.path.splitext(file_name)[0] + ".json"
                )
                
                # Build the command
                command = base_command.copy()
                command.extend(["-o", output_file, docx_file])
                
                # Run pandoc
                result = subprocess.run(
                    command, 
                    capture_output=True, 
                    text=True
                )
                
                if result.returncode != 0:
                    self.update_status(f"Error converting {file_name}: {result.stderr}", True)
                    continue
                
                # Update progress
                self.progress_var.set((i + 1) / total_files * 100)
            
            # All done
            self.update_status(f"Conversion complete. Processed {total_files} files.", True)
            
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            self.update_status(f"Error during conversion: {str(e)}", True)
            print(f"Error details: {error_details}")
    
    def update_status(self, message, finished=False):
        """Update status message and UI state."""
        # Use after() to safely update from a non-main thread
        self.after(0, lambda: self._update_status_ui(message, finished))
    
    def _update_status_ui(self, message, finished):
        """Update UI elements from the main thread."""
        self.status_var.set(message)
        
        if finished:
            self.is_running = False
            self.convert_button.config(state="normal")
    
    def on_close(self):
        """Handle dialog close event."""
        # If conversion is running, confirm before closing
        if self.is_running:
            if messagebox.askyesno("Conversion in Progress", 
                                 "Conversion is still running. Are you sure you want to cancel?"):
                self.is_running = False
                # We can't directly stop the thread, but we can let it complete
                # and the UI will no longer update
            else:
                return  # Don't close if user cancels
        
        # Release grab and destroy window
        self.grab_release()
        self.destroy()

class Section:
    """
    Represents a document section with its heading and content.
    """
    def __init__(self, title, safe_title, level, content, start_index, end_index=None):
        self.title = title
        self.safe_title = safe_title
        self.level = level
        self.content = content
        self.start_index = start_index
        self.end_index = end_index


class StyleProcessor:
    """
    Processes and manages document styles, handling paragraph styles.
    """
    
    def __init__(self, doc: Document):
        self.doc = doc
        self.heading_levels = {}
        self._process_styles()
    
    def _process_styles(self):
        for style in self.doc.styles:
            if style.type != WD_STYLE_TYPE.PARAGRAPH:
                continue
                
            # Handle built-in heading styles
            if style.name.startswith('Heading '):
                try:
                    level = int(style.name.split()[-1])
                    self.heading_levels[style.name] = level
                except (ValueError, IndexError):
                    continue
            
            # Handle custom styles based on headings
            elif hasattr(style, 'base_style') and style.base_style:
                base_name = style.base_style.name
                if base_name.startswith('Heading '):
                    try:
                        level = int(base_name.split()[-1])
                        self.heading_levels[style.name] = level
                    except (ValueError, IndexError):
                        continue
    
    def get_heading_level(self, paragraph: Paragraph):
        if not paragraph.style or not paragraph.style.name:
            return None
        return self.heading_levels.get(paragraph.style.name)


class FilenameManager:
    """Handles creation and management of safe filenames."""
    
    @staticmethod
    def sanitize_filename(title: str, max_length: int = 240) -> str:
        # Remove invalid chars
        safe = re.sub(r'[<>:"/\\|?*]', '', title)
        # Replace whitespace with underscore
        safe = re.sub(r'\s+', '_', safe)
        # Remove duplicate dots
        safe = re.sub(r'\.+', '.', safe)
        # Truncate if needed
        safe = safe[:max_length]
        # Remove leading/trailing dots and underscores
        return safe.strip('._')
    
    def ensure_unique(self, filename: str, used_names: set) -> str:
        base = filename
        counter = 1
        
        while filename in used_names:
            name_parts = base.rsplit('.', 1)
            if len(name_parts) > 1:
                filename = f"{name_parts[0]}_{counter}.{name_parts[1]}"
            else:
                filename = f"{base}_{counter}"
            counter += 1
            
        used_names.add(filename)
        return filename


class DocxSplitter:
    """
    Main class for splitting Word documents by heading level.
    """
    
    def __init__(self, input_path, template_path, status_callback=None, progress_callback=None):
        """
        Initialize the splitter with progress reporting.
        
        Args:
            input_path: Path to document to split
            template_path: Path to template document
            status_callback: Optional callback for status updates
            progress_callback: Optional callback for progress updates (0-100)
        """
        self.input_path = Path(input_path)
        self.template_path = Path(template_path)
        self.status_callback = status_callback or (lambda msg: None)
        self.progress_callback = progress_callback or (lambda percent: None)
        
        # Verify template exists
        if not template_path or not Path(template_path).exists():
            raise ValueError(f"Template document not found: {template_path}")
        
        self.doc = docx.Document(input_path)
        self.style_processor = StyleProcessor(self.doc)
        self.filename_manager = FilenameManager()
        self.sections = []
        
        # Add a cancel flag
        self.cancel_requested = False
    
    def cancel(self):
        """Request cancellation of the current operation."""
        print("DocxSplitter.cancel() method called!")  # Debug
        self.cancel_requested = True
        print(f"cancel_requested flag set to {self.cancel_requested}")  # Debug
        self.status_callback("Cancellation requested")
    
    def _clean_document(self, doc: Document, target_level: int) -> Document:
        """
        Clean document by removing higher-level headings and empty headers.
        Preserves the target heading level for splitting.
        """
        paragraphs_to_remove = []
        
        for idx, para in enumerate(doc.paragraphs):
            # Check for cancellation during cleaning
            if idx % 100 == 0 and self.cancel_requested:
                self.status_callback("Operation canceled during document cleaning")
                return doc
                
            level = self.style_processor.get_heading_level(para)
            
            # Check if this is a heading
            if level is not None:
                # Remove if:
                # 1. It's a heading that's higher in hierarchy than our target level
                #    BUT is not our target level itself
                # 2. It's an empty heading (any level, including target level)
                if (level < target_level and level != target_level) or not para.text.strip():
                    paragraphs_to_remove.append(idx)
                    # Mark the next paragraph for removal if it's empty
                    if idx + 1 < len(doc.paragraphs) and not doc.paragraphs[idx + 1].text.strip():
                        paragraphs_to_remove.append(idx + 1)
        
        # Remove paragraphs in reverse order to maintain correct indices
        for idx in sorted(paragraphs_to_remove, reverse=True):
            if idx < len(doc.paragraphs):  # Safety check
                p = doc.paragraphs[idx]._element
                p.getparent().remove(p)
        
        return doc
    
    def parse_sections(self, target_level: int = 3) -> None:
        """Parse document into sections based on heading level."""
        self.status_callback(f"Parsing document sections at heading level {target_level}...")
        
        # Clean the document first
        self.doc = self._clean_document(self.doc, target_level)
        
        current_section = None
        used_titles = set()
        
        for idx, para in enumerate(self.doc.paragraphs):
            # Check for cancellation periodically
            if idx % 50 == 0 and self.cancel_requested:
                self.status_callback("Operation canceled during parsing")
                self.sections = []  # Clear any partial results
                return
                
            level = self.style_processor.get_heading_level(para)
            
            if level == target_level:
                # Close previous section
                if current_section:
                    current_section.end_index = idx - 1
                    
                    # If current section has no content, remove it
                    if not any(p.text.strip() for p in current_section.content):
                        self.sections.remove(current_section)
                
                # Only create new section if heading has content
                if para.text.strip():
                    # Create new section
                    safe_title = self.filename_manager.sanitize_filename(para.text)
                    unique_title = self.filename_manager.ensure_unique(safe_title, used_titles)
                    
                    current_section = Section(
                        title=para.text,
                        safe_title=unique_title,
                        level=level,
                        content=[],
                        start_index=idx
                    )
                    self.sections.append(current_section)
                    
            elif current_section is not None:
                current_section.content.append(para)
        
        # Handle final section
        if current_section:
            current_section.end_index = len(self.doc.paragraphs) - 1
            # Check if final section is empty
            if not any(p.text.strip() for p in current_section.content):
                self.sections.remove(current_section)
            
        if self.cancel_requested:
            self.status_callback("Operation canceled")
            self.sections = []
            return
                
        self.status_callback(f"Found {len(self.sections)} non-empty sections at heading level {target_level}")
    
    def process_document(self, output_dir, target_level: int = 3, create_zip: bool = True):
        """
        Process the document and output files according to specified options.
        """
        # Make sure we have parsed sections
        if not self.sections:
            self.parse_sections(target_level)
            
        # Check if parsing was canceled
        if self.cancel_requested or not self.sections:
            return None
        
        # Create output directory if it doesn't exist
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        if create_zip:
            return self._create_zip_archive(output_dir)
        else:
            return self._save_individual_files(output_dir)
    
    def _create_section_document(self, section: Section) -> Document:
        """
        Create new document from section content using template.
        """
        # Create new document from template
        new_doc = docx.Document(self.template_path)
        
        # Remove any existing paragraphs from the template
        for p in new_doc.paragraphs[::-1]:  # Iterate in reverse for stable removal
            p._element.getparent().remove(p._element)
        
        # Add the heading as document title (will be first content)
        heading = new_doc.add_heading(section.title, level=section.level)
        
        # Copy content while preserving formatting
        for para in section.content:
            # Skip the heading since we've already added it
            if para.text == section.title:
                continue
                
            # Skip empty paragraphs
            if not para.text.strip():
                continue
                
            new_para = new_doc.add_paragraph()
            
            # Copy runs with careful attribute handling
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                
                # Copy core run properties
                self._copy_core_run_properties(run, new_run)
                
                # Copy extended run properties
                self._copy_extended_run_properties(run, new_run)
                
                # Copy font properties
                self._copy_font_properties(run, new_run)
                
                # Copy style if it exists
                if hasattr(run, 'style') and run.style:
                    try:
                        new_run.style = run.style
                    except Exception as e:
                        print(f"Could not copy run style: {e}")
            
            # Copy paragraph style and properties
            self._copy_paragraph_properties(para, new_para)
                
        return new_doc
    
    def _copy_core_run_properties(self, source_run, target_run) -> None:
        """Copy the core run properties that are guaranteed to exist in python-docx."""
        core_properties = ['bold', 'italic', 'underline']
        
        for prop in core_properties:
            try:
                setattr(target_run, prop, getattr(source_run, prop))
            except Exception as e:
                print(f"Could not copy core property {prop}: {e}")
    
    def _copy_extended_run_properties(self, source_run, target_run) -> None:
        """Copy extended run properties with validation."""
        extended_properties = [
            'all_caps', 'double_strike', 'emboss', 'imprint',
            'outline', 'shadow', 'small_caps', 'strike',
            'subscript', 'superscript'
        ]
        
        for prop in extended_properties:
            try:
                if hasattr(source_run, prop):
                    setattr(target_run, prop, getattr(source_run, prop))
            except Exception as e:
                print(f"Could not copy extended property {prop}: {e}")
    
    def _copy_font_properties(self, source_run, target_run) -> None:
        """Copy font properties with comprehensive handling of colors and highlighting."""
        if not hasattr(source_run, 'font') or not hasattr(target_run, 'font'):
            return
            
        # Copy basic font properties
        if source_run.font.name:
            target_run.font.name = source_run.font.name
            
        if source_run.font.size:
            target_run.font.size = source_run.font.size
        
        # Handle text color (foreground)
        try:
            if hasattr(source_run.font, 'color'):
                if source_run.font.color.rgb is not None:
                    target_run.font.color.rgb = source_run.font.color.rgb
                elif hasattr(source_run.font.color, 'theme_color'):
                    target_run.font.color.theme_color = source_run.font.color.theme_color
        except Exception as e:
            print(f"Could not copy font color: {e}")
            
        # Handle highlighting (background)
        try:
            if hasattr(source_run.font, 'highlight_color'):
                if source_run.font.highlight_color:
                    target_run.font.highlight_color = source_run.font.highlight_color
        except Exception as e:
            print(f"Could not copy highlight color: {e}")
    
    def _copy_paragraph_properties(self, source_para, target_para) -> None:
        """Copy paragraph style, properties, and shading with validation."""
        # Copy style if it exists
        if source_para.style:
            try:
                target_para.style = source_para.style
            except Exception as e:
                print(f"Could not copy paragraph style: {e}")
        
        # Handle paragraph shading/background
        try:
            if hasattr(source_para._element, 'pPr'):
                source_pPr = source_para._element.pPr
                if hasattr(source_pPr, 'shd'):
                    shading = source_pPr.shd
                    if shading is not None and hasattr(target_para._element, 'pPr'):
                        # Ensure pPr exists in target
                        if target_para._element.pPr is None:
                            target_para._element.get_or_add_pPr()
                        # Copy shading element
                        target_para._element.pPr.shd = shading
        except Exception as e:
            print(f"Could not copy paragraph shading: {e}")
        
        # Copy paragraph format properties if they exist
        if hasattr(source_para, 'paragraph_format') and hasattr(target_para, 'paragraph_format'):
            format_properties = [
                'alignment', 'first_line_indent', 'keep_together',
                'keep_with_next', 'left_indent', 'line_spacing',
                'right_indent', 'space_after', 'space_before'
            ]
            
            for prop in format_properties:
                try:
                    source_value = getattr(source_para.paragraph_format, prop)
                    if source_value is not None:
                        setattr(target_para.paragraph_format, prop, source_value)
                except Exception as e:
                    print(f"Could not copy paragraph format property {prop}: {e}")
    
    def _create_zip_archive(self, output_dir: Path):
        """Create zip archive with section documents."""
        zip_path = output_dir / f"{self.input_path.stem}_sections.zip"
        
        # Use temporary directory for intermediate files
        with TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as archive:
                total_sections = len(self.sections)
                for idx, section in enumerate(self.sections, 1):
                    # Explicitly check for cancellation
                    print(f"Checking cancel_requested flag: {self.cancel_requested}")  # Debug
                    if self.cancel_requested:
                        print("Cancellation detected during zip creation!")  # Debug
                        self.status_callback("Operation canceled by user")
                        return None

                    try:
                        # Create document for section
                        doc = self._create_section_document(section)
                        
                        # Save to temp file
                        temp_file = temp_path / f"{section.safe_title}.docx"
                        doc.save(temp_file)
                        
                        # Add to archive
                        archive.write(temp_file, temp_file.name)
                        
                        # Report progress percentage and status
                        percent_complete = int((idx / total_sections) * 100)
                        self.progress_callback(percent_complete)
                        self.status_callback(f"Processed section {idx}/{total_sections}: {section.safe_title}")
                        
                    except Exception as e:
                        self.status_callback(f"Error processing section '{section.safe_title}': {str(e)}")
                        continue
                
                if self.cancel_requested:
                    self.status_callback("Operation canceled while creating archive")
                    return None
                    
            self.status_callback(f"Created archive at: {zip_path}")
            return zip_path
    
    def _save_individual_files(self, output_dir: Path):
        """Save individual document files without zipping."""
        total_sections = len(self.sections)
        for idx, section in enumerate(self.sections, 1):
            # Explicitly check for cancellation before each file
            print(f"Checking cancel_requested flag: {self.cancel_requested}")  # Debug
            if self.cancel_requested:
                print("Cancellation detected during file creation!")  # Debug
                self.status_callback("Operation canceled by user")
                return None
                
            try:
                # Create document for section
                doc = self._create_section_document(section)
                
                # Save to output directory
                output_file = output_dir / f"{section.safe_title}.docx"
                doc.save(output_file)
                
                # Report progress percentage and status
                percent_complete = int((idx / total_sections) * 100)
                self.progress_callback(percent_complete)
                self.status_callback(f"Processed section {idx}/{total_sections}: {section.safe_title}")
                
            except Exception as e:
                self.status_callback(f"Error processing section '{section.safe_title}': {str(e)}")
                continue
            
        if self.cancel_requested:
            self.status_callback("Operation canceled while saving files")
            return None
                
        self.status_callback(f"Saved {total_sections} documents to: {output_dir}")
        return output_dir

class HelpDialog(tk.Toplevel):
    """
    Comprehensive help dialog with tabbed interface covering all application functionality.
    """
    def __init__(self, parent):
        super().__init__(parent)
        self.parent = parent
        
        # Set dialog properties
        self.title("Block Search Help")
        self.geometry("850x650")
        self.minsize(750, 600)
        
        # Make dialog modal
        self.transient(parent)
        self.grab_set()
        
        # Center on parent
        self.update_idletasks()
        parent_x = parent.winfo_x()
        parent_y = parent.winfo_y()
        parent_width = parent.winfo_width()
        parent_height = parent.winfo_height()
        
        width = self.winfo_width()
        height = self.winfo_height()
        
        x = parent_x + (parent_width - width) // 2
        y = parent_y + (parent_height - height) // 2
        
        self.geometry(f"+{x}+{y}")
        
        # Apply simple high-contrast colors
        self.configure(background="#F0F0F0")  # Light gray background
        
        # Store text widgets for resizing
        self.text_widgets = []
        
        # Set up UI
        self.setup_ui()
        
        # Calculate initial wrap width and apply to all text widgets
        self.update_idletasks()  # Ensure sizes are updated
        self.update_wrap_widths()
        
        # Bind resize event
        self.bind("<Configure>", self.on_resize)
        
    def setup_ui(self):
        """Create the tabbed interface and help content."""
        # Use consistent high-contrast colors
        bg_color = "#F0F0F0"
        fg_color = "#000000"
        heading_color = "#000000"
        text_color = "#333333"
        
        # Main frame with padding
        main_frame = tk.Frame(self, bg=bg_color, padx=20, pady=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Create notebook (tabbed interface)
        self.notebook = ttk.Notebook(main_frame)
        self.notebook.pack(fill=tk.BOTH, expand=True)
        
        # Create tabs to match Windows version
        self.create_general_tab(bg_color, fg_color, heading_color, text_color)
        self.create_search_tab(bg_color, fg_color, heading_color, text_color)
        self.create_document_tab(bg_color, fg_color, heading_color, text_color)
        self.create_splitting_tab(bg_color, fg_color, heading_color, text_color)
        self.create_shortcuts_tab(bg_color, fg_color, heading_color, text_color)
        
        # Close button at the bottom
        button_frame = tk.Frame(main_frame, bg=bg_color)
        button_frame.pack(fill=tk.X, pady=(15, 0))
        
        close_button = tk.Button(
            button_frame, 
            text="Close", 
            command=self.destroy,
            bg="#E0E0E0",
            fg="#000000",
            font=("Helvetica", 12)
        )
        close_button.pack(side=tk.RIGHT)
    
    def create_scrollable_frame(self, parent, bg_color):
        """Create a scrollable frame to contain help content."""
        # Container frame
        container = tk.Frame(parent, bg=bg_color)
        container.pack(fill=tk.BOTH, expand=True)
        
        # Add canvas with scrollbar
        canvas = tk.Canvas(container, bg=bg_color, highlightthickness=0)
        scrollbar = ttk.Scrollbar(container, orient="vertical", command=canvas.yview)
        
        # Configure scrollable frame
        scrollable_frame = tk.Frame(canvas, bg=bg_color)
        
        # Create window inside canvas
        frame_id = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        
        # Define resize functions
        def _configure_canvas_scroll(event):
            """Update the scrollregion when the frame size changes."""
            canvas.configure(scrollregion=canvas.bbox("all"))
        
        def _resize_frame_width(event):
            """Resize the inner frame width when canvas changes."""
            # The -20 accounts for some padding to prevent horizontal scrollbar
            canvas_width = event.width - 20 if event.width > 20 else 0
            canvas.itemconfig(frame_id, width=canvas_width)
        
        # Bind frame configuration to adjust scrollbar
        scrollable_frame.bind("<Configure>", _configure_canvas_scroll)
        
        # Bind canvas resize to adjust inner frame width
        canvas.bind("<Configure>", _resize_frame_width)
        
        # Configure canvas scrolling command
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Pack widgets
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Add mouse wheel scrolling
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        
        canvas.bind_all("<MouseWheel>", _on_mousewheel)
        
        # Return the frame where content should be added
        return scrollable_frame
    
    def on_resize(self, event):
        """Update text wrapping when window is resized."""
        # Only process if it's the main window that's resized
        if event.widget == self:
            self.update_wrap_widths()

    def update_wrap_widths(self):
        """Update all text widget wrapping based on current window size."""
        # Calculate available width (account for scrollbar and padding)
        # The -120 accounts for padding, margins, and scrollbar width
        available_width = self.winfo_width() - 200
        
        for text_widget in self.text_widgets:
            if isinstance(text_widget, tk.Label):
                text_widget.config(wraplength=available_width)

    def add_section(self, parent, title, content, bg_color, heading_color, text_color):
        """Add a titled section with content to the help page."""
        # Section frame
        section = tk.Frame(parent, bg=bg_color)
        section.pack(fill=tk.X, pady=(15, 10), padx=10, anchor=tk.W)
        
        # Heading
        if title:
            heading = tk.Label(
                section,
                text=title,
                font=("Helvetica", 14, "bold"),
                bg=bg_color,
                fg=heading_color,
                anchor=tk.W,
                justify=tk.LEFT
            )
            heading.pack(fill=tk.X, pady=(0, 5))
        
        # Default initial wrap width (will be updated on resize)
        initial_wrap_width = 600
        
        # Content - handling multi-line text
        if isinstance(content, str):
            text = tk.Label(
                section,
                text=content,
                font=("Helvetica", 12),
                bg=bg_color,
                fg=text_color,
                justify=tk.LEFT,
                wraplength=initial_wrap_width,
                anchor=tk.W
            )
            text.pack(fill=tk.X, padx=15)
            # Add to list for resizing
            self.text_widgets.append(text)
            
        elif isinstance(content, list):
            # For bullet points or step lists
            for item in content:
                # Check if item starts with a number followed by period (numbered list)
                if re.match(r'^\d+\.', item):
                    # It's a numbered item, preserve the format
                    bullet_text = item
                    indent = 15
                else:
                    # It's a regular bullet item
                    bullet = "•"
                    bullet_text = f"{bullet} {item}"
                    indent = 25
                    
                    # Check if it's an indented bullet (starts with spaces)
                    if item.startswith('   '):
                        indent = 40
                
                point = tk.Label(
                    section,
                    text=bullet_text,
                    font=("Helvetica", 12),
                    bg=bg_color,
                    fg=text_color,
                    justify=tk.LEFT,
                    wraplength=initial_wrap_width - indent,  # Narrower for bullets/numbers
                    anchor=tk.W
                )
                point.pack(fill=tk.X, padx=indent, pady=(2, 2), anchor=tk.W)
                # Add to list for resizing
                self.text_widgets.append(point)
    
    def create_general_tab(self, bg_color, fg_color, heading_color, text_color):
        """Create the general help tab."""
        tab = tk.Frame(self.notebook, bg=bg_color)
        self.notebook.add(tab, text="General")
        
        content_frame = self.create_scrollable_frame(tab, bg_color)
        
        # Add introductory section
        intro_text = (
            "Block Search helps you quickly search for and convert JSON files to Word documents."
        )
        self.add_section(content_frame, "Block Search Utility", intro_text, bg_color, heading_color, text_color)
        
        # Add getting started section
        getting_started = [
            "1. Select the folder with your JSON files from File → Browse Directory",
            "2. Optionally, select a template document for conversion",
            "3. Type in the search box to find documents",
            "4. Click on a document to convert and copy its contents"
        ]
        self.add_section(content_frame, "Getting Started:", getting_started, bg_color, heading_color, text_color)

    def create_search_tab(self, bg_color, fg_color, heading_color, text_color):
        """Create the search tab."""
        tab = tk.Frame(self.notebook, bg=bg_color)
        self.notebook.add(tab, text="Searching your Blocks")
        
        content_frame = self.create_scrollable_frame(tab, bg_color)
        
        # Add basic search explanation
        basic_search = (
            "Type search terms in the search box to find matching documents. "
            "Results update as you type."
        )
        self.add_section(content_frame, "Basic Search:", basic_search, bg_color, heading_color, text_color)
        
        # Add search tips
        search_tips = [
            "Press Cmd+F or Ctrl+Shift+Space to focus the search field",
            "Type any part of a filename to filter the list",
            "Spaces in search terms are treated as AND operators",
            "Search is performed across filenames only",
            "Send the doc by double clicking it or selecting it and pressing Enter."
        ]
        self.add_section(content_frame, "Search Tips:", search_tips, bg_color, heading_color, text_color)
        
        # Add sorting explanation
        sorting_text = (
            "You can sort the file list by clicking on column headers or using the sort buttons. "
            "To change sort direction, click the button for your current sort, click the sort header, or toggle the 'Reverse' checkbox."
        )
        self.add_section(content_frame, "Sorting Files:", sorting_text, bg_color, heading_color, text_color)
        
        # Add sort options
        sort_options = [
            "Name: Sort alphabetically by filename",
            "Path: Sort by relative folder path",
            "Last Opened: Sort by most recently accessed files"
        ]
        self.add_section(content_frame, "Sort Options:", sort_options, bg_color, heading_color, text_color)

    def create_document_tab(self, bg_color, fg_color, heading_color, text_color):
        """Create the document operations tab."""
        tab = tk.Frame(self.notebook, bg=bg_color)
        self.notebook.add(tab, text="Sending to Doc")
        
        content_frame = self.create_scrollable_frame(tab, bg_color)
        
        # Add document transfer section
        doc_transfer = (
            "When you select a document from search results, the content is handled based on your settings:"
        )
        
        transfer_options = [
            "Clipboard Mode: Content is copied to clipboard (default when no target is set)",
            "Target Document Mode: Content is pasted into your selected open Word document"
        ]
        
        self.add_section(content_frame, "Document Content Transfer:", doc_transfer, bg_color, heading_color, text_color)
        self.add_section(content_frame, "", transfer_options, bg_color, heading_color, text_color)
        
        # Add using open Word documents section
        word_docs = (
            "The Target dropdown shows currently open Word documents:"
        )
        
        word_docs_steps = [
            "1. Select a document from the dropdown to set it as the target",
            "2. Click the refresh button (⟳) to update the list of open documents",
            "3. Double-click a file or press Enter to send that block to your speech doc"
        ]
        
        self.add_section(content_frame, "Sending to Open Speech Docs:", word_docs, bg_color, heading_color, text_color)
        self.add_section(content_frame, "", word_docs_steps, bg_color, heading_color, text_color)
        
        # Add template section
        template_text = (
            "Templates allow you to maintain consistent formatting when converting JSON to Word. "
            "The application will look for 'template_doc.docx' in your selected directory by default, "
            "or you can browse to select a different template. Your 'template_doc.docx' should be a blank"
            "Verbatimized document."
        )
        self.add_section(content_frame, "Using Templates:", template_text, bg_color, heading_color, text_color)

    def create_splitting_tab(self, bg_color, fg_color, heading_color, text_color):
        """Create the document splitter tab."""
        tab = tk.Frame(self.notebook, bg=bg_color)
        self.notebook.add(tab, text="Document Splitter")
        
        content_frame = self.create_scrollable_frame(tab, bg_color)
        
        # Add overview section
        overview = (
            "The Document Splitter allows you to break a large Word document into smaller files based on heading levels. "
            "This is useful for extracting blocks from a large file."
        )
        self.add_section(content_frame, "Overview:", overview, bg_color, heading_color, text_color)
        
        # Add how to access section
        access_text = (
            "Open the splitter from Document Tools → Split Documents by Headings..."
        )
        self.add_section(content_frame, "How to Access:", access_text, bg_color, heading_color, text_color)
        
        # Add usage steps
        usage_steps = (
            "1. Select Input Document: Choose the Word document you want to split\n\n"
            "2. Choose Template Document: Select a document to use as a template for output files. "
            "This should be a blank, Verbatimized Word document.\n\n"
            "3. Select Heading Level: Choose which heading level to split at (Heading 1-4). "
            "The document will be divided at each heading of the selected level. "
            "Each heading and its content will become a separate document.\n\n"
            "4. Choose Output Options: Create ZIP Archive - Package all split documents into a single ZIP file, or "
            "Individual Files - Save each section as a separate document.\n\n"
            "5. Select Output Location: Choose where to save the output files\n\n"
            "6. Process Document: Click the button to start the splitting process"
        )
        self.add_section(content_frame, "Using the Splitter:", usage_steps, bg_color, heading_color, text_color)
        
        # Add tips section
        tips = [
            "Results will be best for files that use consistent heading styles",
            "Keep searchability in mind; your heading titles will become your doc titles, which will become the names of the blocks you are searching for.",
            "Use the ZIP option for easier file sharing when creating many documents",
            "For best results, use a completely blank Verbatim document as your template"
        ]
        self.add_section(content_frame, "Tips:", tips, bg_color, heading_color, text_color)

    def create_shortcuts_tab(self, bg_color, fg_color, heading_color, text_color):
        """Create the keyboard shortcuts tab."""
        tab = tk.Frame(self.notebook, bg=bg_color)
        self.notebook.add(tab, text="Keyboard Shortcuts")
        
        content_frame = self.create_scrollable_frame(tab, bg_color)
        
        # Application shortcuts section
        app_shortcuts_header = "Application Shortcuts:"
        if sys.platform == 'darwin':  # macOS
            app_shortcuts = [
                "⌘F: Focus the search field",
                "⌃⇧Space: Alternate shortcut to focus search",
                "⌘R: Refresh the list of open Word documents",
                "⌘Q: Quit application",
            ]
        else:  # Windows/Linux
            app_shortcuts = [
                "Ctrl+F: Focus the search field",
                "Ctrl+Shift+Space: Alternate shortcut to focus search",
                "Ctrl+R: Refresh the list of open Word documents",
                "Alt+F4, Ctrl+Q: Quit application",
                "F1: Show this help dialog"
            ]
        
        self.add_section(content_frame, app_shortcuts_header, app_shortcuts, bg_color, heading_color, text_color)
        
        # Search result navigation section
        nav_shortcuts_header = "Search Result Navigation:"
        nav_shortcuts = [
            "Up/Down: Navigate through search results",
            "Enter: Select document (process the selected file)",
            "Double-click: Select document (same as Enter)"
        ]
        
        self.add_section(content_frame, nav_shortcuts_header, nav_shortcuts, bg_color, heading_color, text_color)

    def open_help_dialog(self):
        """Open the comprehensive help dialog."""
        HelpDialog(self.root)

    def open_shortcuts_help(self):
        """Open the help dialog directly to the shortcuts tab."""
        help_dialog = HelpDialog(self.root)
        help_dialog.notebook.select(4)  # Index of the shortcuts tab (0-based)

class DocSplitterDialog(tk.Toplevel):
    """
    Dialog for splitting Word documents by heading level.
    """
    
    def __init__(self, parent):
        super().__init__(parent)
        self.parent = parent
        
        # Set dialog properties
        self.title("Split Document by Headings")
        self.geometry("650x450")
        self.resizable(True, True)
        self.minsize(600, 400)
        
        # Mac-specific dialog setup
        self.transient(parent)  # Dialog follows parent window
        self.grab_set()  # Make dialog modal
        
        # Center on parent
        self.update_idletasks()
        parent_x = parent.winfo_x()
        parent_y = parent.winfo_y()
        parent_width = parent.winfo_width()
        parent_height = parent.winfo_height()
        
        width = self.winfo_width()
        height = self.winfo_height()
        
        x = parent_x + (parent_width - width) // 2
        y = parent_y + (parent_height - height) // 2
        
        self.geometry(f"+{x}+{y}")
        
        # Initialize variables
        self.input_path = None
        self.template_path = None
        self.output_dir = None
        self.cancel_requested = False  # Flag for cancellation
        self.process_thread = None     # Reference to processing thread
        
        # Set up the UI
        self.setup_ui()
        
        # Set up the protocol to handle dialog close
        self.protocol("WM_DELETE_WINDOW", self.on_close)
        
    def setup_ui(self):
        """Create the user interface."""
        # Main frame with padding
        main_frame = tk.Frame(self, bg="#F0F0F0", padx=20, pady=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Description label at the top
        description_text = (
            "Split a Word document into multiple files based on heading levels. "
            "This tool divides the document at each heading of the specified level "
            "and creates a separate file for each section.\n\n"
            "TIP: For best results, use a completely blank Verbatim document as your template."
        )
        
        description_label = tk.Label(
            main_frame,
            text=description_text,
            bg="#F0F0F0",
            fg="#000000",
            font=("Helvetica", 12),
            justify=tk.LEFT,
            wraplength=600
        )
        description_label.pack(fill=tk.X, pady=(0, 15))
        
        # Input document selection
        input_frame = tk.Frame(main_frame, bg="#F0F0F0")
        input_frame.pack(fill=tk.X, pady=5)
        
        input_label = tk.Label(
            input_frame,
            text="Input Document:",
            bg="#F0F0F0",
            fg="#000000",
            font=("Helvetica", 12),
            width=15,
            anchor=tk.W
        )
        input_label.pack(side=tk.LEFT)
        
        self.input_field = tk.Entry(
            input_frame,
            bg="#FFFFFF",
            fg="#000000",
            font=("Helvetica", 12),
            readonlybackground="#FFFFFF"
        )
        self.input_field.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(5, 5))
        self.input_field.config(state="readonly")
        
        input_button = tk.Button(
            input_frame,
            text="Browse...",
            bg="#E0E0E0",
            fg="#000000",
            font=("Helvetica", 12),
            command=self.browse_input_document
        )
        input_button.pack(side=tk.LEFT, padx=(5, 0))
        
        # Template document selection
        template_frame = tk.Frame(main_frame, bg="#F0F0F0")
        template_frame.pack(fill=tk.X, pady=5)
        
        template_label = tk.Label(
            template_frame,
            text="Template Doc:",
            bg="#F0F0F0",
            fg="#000000",
            font=("Helvetica", 12),
            width=15,
            anchor=tk.W
        )
        template_label.pack(side=tk.LEFT)
        
        self.template_field = tk.Entry(
            template_frame,
            bg="#FFFFFF",
            fg="#000000",
            font=("Helvetica", 12),
            readonlybackground="#FFFFFF"
        )
        self.template_field.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(5, 5))
        self.template_field.config(state="readonly")
        
        template_button = tk.Button(
            template_frame,
            text="Browse...",
            bg="#E0E0E0",
            fg="#000000",
            font=("Helvetica", 12),
            command=self.browse_template_document
        )
        template_button.pack(side=tk.LEFT, padx=(5, 0))
        
        # Heading level selection
        level_frame = tk.Frame(main_frame, bg="#F0F0F0")
        level_frame.pack(fill=tk.X, pady=5)
        
        level_label = tk.Label(
            level_frame,
            text="Split at Level:",
            bg="#F0F0F0",
            fg="#000000",
            font=("Helvetica", 12),
            width=15,
            anchor=tk.W
        )
        level_label.pack(side=tk.LEFT)
        
        self.level_var = tk.StringVar(value="Heading 3")
        self.level_combo = ttk.Combobox(
            level_frame,
            textvariable=self.level_var,
            values=["Heading 1", "Heading 2", "Heading 3", "Heading 4"],
            font=("Helvetica", 12),
            state="readonly",
            width=20
        )
        self.level_combo.pack(side=tk.LEFT, padx=(5, 0))
        
        # Output options
        options_frame = tk.Frame(main_frame, bg="#F0F0F0")
        options_frame.pack(fill=tk.X, pady=10)
        
        self.zip_var = tk.BooleanVar(value=True)
        zip_check = tk.Checkbutton(
            options_frame,
            text="Create ZIP archive of documents",
            variable=self.zip_var,
            bg="#F0F0F0",
            fg="#000000",
            font=("Helvetica", 12),
            selectcolor="#FFFFFF"
        )
        zip_check.pack(side=tk.LEFT, padx=(15, 0))
        
        # Output directory selection
        output_frame = tk.Frame(main_frame, bg="#F0F0F0")
        output_frame.pack(fill=tk.X, pady=5)
        
        output_label = tk.Label(
            output_frame,
            text="Output Location:",
            bg="#F0F0F0",
            fg="#000000",
            font=("Helvetica", 12),
            width=15,
            anchor=tk.W
        )
        output_label.pack(side=tk.LEFT)
        
        self.output_field = tk.Entry(
            output_frame,
            bg="#FFFFFF",
            fg="#000000",
            font=("Helvetica", 12),
            readonlybackground="#FFFFFF"
        )
        self.output_field.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(5, 5))
        self.output_field.config(state="readonly")
        
        output_button = tk.Button(
            output_frame,
            text="Browse...",
            bg="#E0E0E0",
            fg="#000000",
            font=("Helvetica", 12),
            command=self.browse_output_directory
        )
        output_button.pack(side=tk.LEFT, padx=(5, 0))
        
        # Status and progress section
        status_frame = tk.Frame(main_frame, bg="#F0F0F0")
        status_frame.pack(fill=tk.X, pady=(15, 5))
        
        self.status_var = tk.StringVar(value="Ready")
        status_label = tk.Label(
            status_frame,
            textvariable=self.status_var,
            bg="#F0F0F0",
            fg="#505050",
            font=("Helvetica", 11),
            anchor=tk.W
        )
        status_label.pack(fill=tk.X)
        
        self.progress_bar = ttk.Progressbar(
            status_frame,
            orient="horizontal",
            mode="determinate",
            length=100
        )
        self.progress_bar.pack(fill=tk.X, pady=(5, 0))
        self.progress_bar.pack_forget()  # Hide initially
        
        # Buttons at bottom
        button_frame = tk.Frame(main_frame, bg="#F0F0F0")
        button_frame.pack(fill=tk.X, pady=(15, 0), side=tk.BOTTOM)
        
        self.process_button = tk.Button(
            button_frame,
            text="Process Document",
            bg="#E0E0E0",
            fg="#000000",
            font=("Helvetica", 12),
            command=self.process_document,
            padx=10
        )
        self.process_button.pack(side=tk.RIGHT, padx=5)
        
        self.cancel_button = tk.Button(
            button_frame,
            text="Cancel",
            bg="#E0E0E0",
            fg="#000000",
            font=("Helvetica", 12),
            command=self.on_close,
            padx=10
        )
        self.cancel_button.pack(side=tk.RIGHT, padx=5)
        
    def browse_input_document(self):
        """Open file dialog to select input document."""
        file_path = filedialog.askopenfilename(
            title="Select Document to Split",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        
        if file_path:
            self.input_path = file_path
            self.input_field.config(state="normal")
            self.input_field.delete(0, tk.END)
            self.input_field.insert(0, file_path)
            self.input_field.config(state="readonly")
            
            # Set default output directory to input file's directory
            if not self.output_dir:
                self.output_dir = os.path.dirname(file_path)
                self.output_field.config(state="normal")
                self.output_field.delete(0, tk.END)
                self.output_field.insert(0, self.output_dir)
                self.output_field.config(state="readonly")
                
    def browse_template_document(self):
        """Open file dialog to select template document."""
        file_path = filedialog.askopenfilename(
            title="Select Template Document",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        
        if file_path:
            self.template_path = file_path
            self.template_field.config(state="normal")
            self.template_field.delete(0, tk.END)
            self.template_field.insert(0, file_path)
            self.template_field.config(state="readonly")
                
    def browse_output_directory(self):
        """Open directory dialog to select output location."""
        directory = filedialog.askdirectory(
            title="Select Output Directory"
        )
        
        if directory:
            self.output_dir = directory
            self.output_field.config(state="normal")
            self.output_field.delete(0, tk.END)
            self.output_field.insert(0, directory)
            self.output_field.config(state="readonly")
    
    def update_status(self, message):
        """Update status display with thread safety."""
        self.status_var.set(message)
        self.update_idletasks()  # Force UI update
        
    def update_progress(self, percent):
        """Update progress bar with thread safety."""
        self.progress_bar['value'] = percent
        self.update_idletasks()  # Force UI update
        
    def validate_inputs(self):
        """Validate that all required inputs are provided."""
        if not self.input_path:
            messagebox.showerror("Missing Input", "Please select an input document.")
            return False
        
        if not os.path.exists(self.input_path):
            messagebox.showerror("Input Error", "Input document does not exist.")
            return False
        
        if not self.output_dir:
            messagebox.showerror("Missing Output", "Please select an output location.")
            return False
        
        # Create output directory if it doesn't exist
        if not os.path.exists(self.output_dir):
            try:
                os.makedirs(self.output_dir)
            except OSError as e:
                messagebox.showerror("Output Error", f"Could not create output directory: {str(e)}")
                return False
        
        # Check template if specified
        if self.template_path and not os.path.exists(self.template_path):
            messagebox.showerror("Template Error", "Template document does not exist.")
            return False
            
        return True
    
    def get_heading_level(self):
        """Get numeric heading level from combobox selection."""
        heading_text = self.level_var.get()
        try:
            # Extract the number from "Heading X"
            return int(heading_text.split()[-1])
        except (ValueError, IndexError):
            # Default to heading level 3 if there's any problem
            return 3
    
    def process_document(self):
        """Start document processing in a separate thread."""
        if not self.validate_inputs():
            return
            
        # Show progress bar
        self.progress_bar.pack(fill=tk.X, pady=(5, 0))
        self.progress_bar['value'] = 0
        
        # Disable process button during processing
        self.process_button.config(state="disabled")
        
        # Clear cancellation flag
        self.cancel_requested = False
        
        # Handle template creation if needed
        if not self.template_path:
            self.update_status("Creating default template...")
            
            try:
                # Create a simple empty document as template
                default_doc = docx.Document()
                
                with TemporaryDirectory() as temp_dir:
                    temp_path = os.path.join(temp_dir, "default_template.docx")
                    default_doc.save(temp_path)
                    self.template_path = temp_path
                    
                    # Continue with processing using temp template
                    self._start_processing_thread()
            except Exception as e:
                messagebox.showerror("Template Error", f"Failed to create default template: {str(e)}")
                self.process_button.config(state="normal")
                self.progress_bar.pack_forget()
        else:
            # Use existing template
            self._start_processing_thread()
    
    def _start_processing_thread(self):
        """Start the document processing in a separate thread."""
        self.process_thread = threading.Thread(target=self._process_document_thread)
        self.process_thread.daemon = True
        self.process_thread.start()
    
    def _process_document_thread(self):
        """Process document in a separate thread to keep UI responsive."""
        try:
            # Get the heading level
            heading_level = self.get_heading_level()
            create_zip = self.zip_var.get()
            
            # Create splitter with callbacks
            splitter = DocxSplitter(
                self.input_path,
                self.template_path,
                status_callback=self.update_status,
                progress_callback=self.update_progress
            )
            
            # Process the document
            result_path = splitter.process_document(
                self.output_dir,
                target_level=heading_level,
                create_zip=create_zip
            )
            
            # Check if operation was canceled
            if result_path is None:
                self.after(0, self._cleanup_canceled_operation)
                return
                
            # Ensure progress bar shows 100%
            self.after(0, lambda: self.update_progress(100))
            
            # Show success message
            section_count = len(splitter.sections)
            self.after(0, lambda: self._show_completion_message(result_path, section_count, heading_level))
            
        except Exception as e:
            # Show error in main thread
            self.after(0, lambda: self._show_error(str(e)))
    
    def _cleanup_canceled_operation(self):
        """Clean up UI after canceled operation."""
        self.update_status("Operation canceled")
        self.process_button.config(state="normal")
        self.progress_bar.pack_forget()
    
    def _show_completion_message(self, result_path, section_count, heading_level):
        """Show completion message and reset UI."""
        self.update_status(f"Processing complete. Created {section_count} documents.")
        self.process_button.config(state="normal")
        
        # Show success message
        msg_text = (
            f"Document processed successfully.\n"
            f"Found {section_count} sections at heading level {heading_level}.\n\n"
            f"Output saved to: {result_path}"
        )
        
        messagebox.showinfo("Success", msg_text)
        
        # Reset for next operation
        self.input_path = None
        self.input_field.config(state="normal")
        self.input_field.delete(0, tk.END)
        self.input_field.config(state="readonly")
        self.progress_bar.pack_forget()
    
    def _show_error(self, error_message):
        """Show error message and reset UI."""
        self.update_status("Error during processing")
        self.process_button.config(state="normal")
        self.progress_bar.pack_forget()
        
        messagebox.showerror("Error", f"Failed to process document:\n{error_message}")
    
    def on_close(self):
        """Handle dialog close with cancellation support."""
        # Check if processing is active
        if self.process_thread and self.process_thread.is_alive() and not self.cancel_requested:
            # Ask user to confirm cancellation
            if messagebox.askyesno(
                "Cancel Operation",
                "A document split operation is in progress.\nDo you want to cancel it?"
            ):
                # Set cancellation flag
                self.cancel_requested = True
                self.update_status("Canceling operation...")
                
                # Don't close yet - wait for thread to acknowledge cancellation
                # We'll check again in 500ms
                self.after(500, self.on_close)
            return
        
        # If thread is done or cancellation is in progress, close the dialog
        self.grab_release()
        self.destroy()

def main():
    # Check for Pandoc
    if not check_pandoc():
        messagebox.showerror("Missing Dependency", 
                            "Pandoc is required but not found.\n\n"
                            "Please install Pandoc using Homebrew:\n"
                            "brew install pandoc\n\n"
                            "Or download from: https://pandoc.org/installing.html")
        sys.exit(1)
    
    # Create main window
    root = tk.Tk()
    root.title("JSON to Word")
    root.geometry("1000x700")  # Increased window size to accommodate wider columns
    
    # Set macOS appearance
    root.tk.call('tk', 'scaling', 2.0)  # Improve HiDPI display
    
    # Create application
    app = JSONToWordConverter(root)
    
    # Center window on screen
    root.eval('tk::PlaceWindow . center')
    
    # Start the application
    root.mainloop()

if __name__ == "__main__":
    main()